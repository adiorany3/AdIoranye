"""Power features for Adioranye AI.

Production-oriented extension layer for model optimization, persistent memory,
lightweight RAG, response cache, usage/cost logging, per-intent prompt templates,
model benchmarking, adaptive model scoring, and circuit breaker protection.

Designed for Streamlit Cloud: SQLite only, no heavy dependencies required.
"""

from __future__ import annotations

import csv
import hashlib
import io
import json
import math
import re
import sqlite3
import time
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from zoneinfo import ZoneInfo

from ai_core import call_api_once, model_cost_tier, model_price

try:
    from db_guard import ensure_database_ready, maybe_create_periodic_backup, default_backup_dir, default_max_backups
except Exception:  # pragma: no cover - DB guard is optional fallback
    ensure_database_ready = None  # type: ignore
    maybe_create_periodic_backup = None  # type: ignore
    default_backup_dir = lambda: '.db_backups'  # type: ignore
    default_max_backups = lambda: 10  # type: ignore

try:
    from critical_current_layer import (
        build_critical_answer_instruction,
        calculate_freshness_score,
        datetime_to_ts as _critical_datetime_to_ts,
        detect_critical_question,
        format_issue_report,
    )
except Exception:  # keep older deployments alive if the helper file is missing
    def detect_critical_question(text: str) -> Dict[str, Any]:
        return {"is_critical": False, "score": 0, "matched_keywords": [], "mode": "normal"}
    def build_critical_answer_instruction(user_text: str, detection: Optional[Dict[str, Any]] = None) -> str:
        return ""
    def calculate_freshness_score(published_at: Any = "", scraped_at: Any = "") -> Dict[str, Any]:
        return {"score": 45.0, "bucket": "unknown", "age_days": None, "published_ts": 0.0}
    def _critical_datetime_to_ts(value: Any) -> float:
        return 0.0
    def format_issue_report(query: str, claims: List[Dict[str, Any]], docs: Optional[List[Dict[str, Any]]] = None) -> str:
        return "Belum ada data KB yang cukup untuk isu: " + str(query)


try:
    from ai_quality_control import (
        build_mode_system_instruction,
        build_quality_footer,
        infer_answer_mode,
        mode_help_text,
        mode_policy,
        normalize_answer_mode,
        parse_mode_command,
        score_answer_quality,
        verify_and_repair_answer,
    )
except Exception:  # keep app alive if the quality layer is absent
    def normalize_answer_mode(value: Any = "auto") -> str:
        return "auto"
    def infer_answer_mode(user_text: str, requested_mode: str = "auto", intent: str = "") -> str:
        return "auto"
    def mode_policy(mode: str) -> Dict[str, Any]:
        return {"temperature_cap": 0.3, "token_multiplier": 1.0, "force_rag": False, "strict_rag": False, "verifier": False, "min_sources": 0}
    def build_mode_system_instruction(mode: str, user_text: str = "") -> str:
        return ""
    def score_answer_quality(**kwargs: Any) -> Any:
        class _Q:
            score = 0.7
            level = "cukup"
            needs_verification = False
            reasons = ["quality_layer_unavailable"]
            metrics = {}
            def to_dict(self) -> Dict[str, Any]:
                return {"score": self.score, "level": self.level, "needs_verification": False, "reasons": self.reasons, "metrics": {}}
        return _Q()
    def build_quality_footer(result: Any, mode: str) -> str:
        return ""
    def verify_and_repair_answer(**kwargs: Any) -> Tuple[str, Dict[str, Any]]:
        return str(kwargs.get("answer") or ""), {"skipped": True}
    def parse_mode_command(text: str) -> Optional[str]:
        return None
    def mode_help_text(current_mode: str = "auto") -> str:
        return "Mode jawaban belum tersedia."


try:
    from performance_optimizer import (
        QueryPlan,
        build_performance_report,
        latency_budget_seconds,
        retrieval_precision_estimate,
        rerank_sources,
        rewrite_query,
        semantic_similarity,
    )
except Exception:  # keep older deployments alive if performance layer is absent
    class QueryPlan:  # type: ignore
        def __init__(self, original_query='', rewritten_query='', extra_terms=None, is_casual=False, should_use_rag=True, reason='fallback'):
            self.original_query = original_query
            self.rewritten_query = rewritten_query or original_query
            self.extra_terms = extra_terms or []
            self.is_casual = is_casual
            self.should_use_rag = should_use_rag
            self.reason = reason
        def to_dict(self):
            return dict(original_query=self.original_query, rewritten_query=self.rewritten_query, extra_terms=self.extra_terms, is_casual=self.is_casual, should_use_rag=self.should_use_rag, reason=self.reason)
    def rewrite_query(user_text: str, intent: str = '', answer_mode: str = 'auto', max_terms: int = 16):
        return QueryPlan(user_text, user_text, [], False, True, 'performance_layer_unavailable')
    def rerank_sources(query: str, sources=None, limit: int = 5, diversity: bool = True):
        return list(sources or [])[:limit]
    def semantic_similarity(a, b):
        return 0.0
    def retrieval_precision_estimate(query: str, sources):
        return {'precision': 0.0, 'avg_similarity': 0.0, 'relevant_count': 0, 'source_count': len(sources or [])}
    def latency_budget_seconds(answer_mode: str = 'auto', intent: str = '', user_text: str = ''):
        return 60
    def build_performance_report(metrics):
        return str(metrics)


try:
    from music_chart_tools import (
        build_music_chart_context,
        build_music_chart_fallback_answer,
        fetch_indonesia_music_charts,
        is_indonesia_music_chart_query,
        is_music_chart_query,
        music_chart_result_to_pseudo_source,
    )
except Exception:  # keep app alive if music chart helper is absent
    def is_music_chart_query(text: Any) -> bool:
        return False
    def is_indonesia_music_chart_query(text: Any) -> bool:
        return False
    def fetch_indonesia_music_charts(limit: int = 10, timeout: int = 8) -> Any:
        class _R:
            ok = False
            source_name = ""
            source_url = ""
            fetched_at_wib = ""
            items = ()
            note = "music_chart_tools unavailable"
            errors = ("music_chart_tools unavailable",)
            def to_dict(self):
                return {"ok": False, "errors": list(self.errors)}
        return _R()
    def build_music_chart_context(result: Any, max_items: int = 10) -> str:
        return ""
    def build_music_chart_fallback_answer(result: Any) -> str:
        return "Data tangga lagu terbaru belum tersedia."
    def music_chart_result_to_pseudo_source(result: Any) -> Dict[str, Any]:
        return {}


try:
    from live_knowledge_fallback import (
        LiveSearchResult,
        build_live_search_system_instruction,
        live_result_to_rag_sources,
        save_live_result_to_kb,
        should_trigger_live_fallback,
        tavily_live_search,
    )
except Exception:  # keep app alive if live fallback helper is absent
    class LiveSearchResult:  # type: ignore
        ok = False
        def to_dict(self):
            return {"ok": False, "reason": "live_knowledge_fallback unavailable"}
    def should_trigger_live_fallback(user_text: str, **kwargs: Any) -> Dict[str, Any]:
        return {"use": False, "reason": "live_knowledge_fallback_unavailable"}
    def tavily_live_search(query: str, **kwargs: Any) -> Any:
        class _R:
            ok = False
            reason = "live_knowledge_fallback_unavailable"
            errors = ["live_knowledge_fallback unavailable"]
            items = []
            def __init__(self, query_text: str = ""):
                self.query = query_text
            def to_dict(self):
                return {"ok": False, "query": self.query, "reason": self.reason, "errors": self.errors}
        return _R(query)
    def live_result_to_rag_sources(result: Any, **kwargs: Any) -> List[Dict[str, Any]]:
        return []
    def save_live_result_to_kb(store: Any, result: Any, **kwargs: Any) -> Dict[str, Any]:
        return {"saved": 0, "skipped": True}
    def build_live_search_system_instruction(result: Any) -> str:
        return ""

try:
    from hallucination_guard import (
        append_guard_note,
        apply_temperature_policy,
        build_guard_system_instruction,
        build_insufficient_evidence_answer,
        detect_high_risk_question,
        evaluate_evidence_gate,
        lightweight_claim_risk,
    )
except Exception:  # keep app alive if the guard file is absent
    def detect_high_risk_question(text: str, intent: str = "") -> Dict[str, Any]:
        return {"is_high_risk": False, "score": 0, "matched_keywords": [], "mode": "normal"}
    def evaluate_evidence_gate(user_text: str, rag_sources: Optional[List[Dict[str, Any]]] = None, **kwargs: Any) -> Any:
        class _FallbackGuard:
            enabled = False
            is_high_risk = False
            allow_answer = True
            reason = "guard_unavailable"
            mode = "normal"
            def to_dict(self) -> Dict[str, Any]:
                return {"enabled": False, "allow_answer": True, "reason": "guard_unavailable"}
        return _FallbackGuard()
    def build_guard_system_instruction(user_text: str, rag_sources: Optional[List[Dict[str, Any]]] = None, guard: Any = None) -> str:
        return ""
    def build_insufficient_evidence_answer(user_text: str, guard: Any, intent: str = "") -> str:
        return "Data belum cukup kuat di Knowledge Base untuk menjawab pertanyaan ini secara aman."
    def apply_temperature_policy(temperature: float, guard: Any = None) -> float:
        return float(temperature or 0.3)
    def append_guard_note(answer: str, rag_sources: Optional[List[Dict[str, Any]]] = None, **kwargs: Any) -> str:
        return answer
    def lightweight_claim_risk(answer: str) -> Dict[str, Any]:
        return {}

WIB_TZ = ZoneInfo("Asia/Jakarta")
DEFAULT_POWER_DB = ".adioranye_power.db"


# =========================
# General helpers
# =========================

def now_wib_text() -> str:
    return datetime.now(WIB_TZ).strftime("%Y-%m-%d %H:%M:%S WIB")


def _utc_ts() -> float:
    return time.time()


def _timestamp_to_wib(ts: float) -> str:
    try:
        if not ts:
            return ""
        return datetime.fromtimestamp(float(ts), WIB_TZ).strftime("%Y-%m-%d %H:%M:%S WIB")
    except Exception:
        return ""


def _safe_json(data: Any) -> str:
    try:
        return json.dumps(data or {}, ensure_ascii=False, default=str)
    except Exception:
        return "{}"


def _safe_json_loads(text: Any) -> Dict[str, Any]:
    try:
        value = json.loads(str(text or "{}"))
        return value if isinstance(value, dict) else {}
    except Exception:
        return {}


def _row_value(row: Any, key: str, index: int = 0, default: Any = None) -> Any:
    """Read a value safely from sqlite3.Row, dict, tuple/list, or None."""
    if row is None:
        return default
    if isinstance(row, dict):
        return row.get(key, default)
    try:
        return row[key]
    except Exception:
        pass
    try:
        return row[index]
    except Exception:
        return default


def _row_to_dict(row: Any, columns: Optional[List[str]] = None) -> Dict[str, Any]:
    if row is None:
        return {}
    if isinstance(row, dict):
        return dict(row)
    try:
        return dict(row)
    except Exception:
        pass
    if columns and isinstance(row, (tuple, list)):
        return {columns[i]: row[i] if i < len(row) else None for i in range(len(columns))}
    return {}


def _tokenize(text: str) -> List[str]:
    words = re.findall(r"[a-zA-Z0-9_\-]{3,}", str(text or "").lower())
    stop = {
        "yang", "dan", "atau", "untuk", "dengan", "dari", "pada", "agar", "jadi", "ini", "itu",
        "dalam", "akan", "bisa", "saya", "anda", "kami", "kamu", "apa", "bagaimana", "karena",
        "the", "and", "for", "with", "from", "that", "this", "are", "was", "were", "you", "your",
    }
    return [w for w in words if w not in stop]


def _score_text(query: str, text: str) -> float:
    q_terms = _tokenize(query)
    if not q_terms:
        return 0.0
    hay = str(text or "").lower()
    tokens = set(_tokenize(text))
    score = 0.0
    for term in q_terms:
        if term in tokens:
            score += 2.0
        if term in hay:
            score += 0.5
    q_clean = " ".join(q_terms[:6])
    if q_clean and q_clean in hay:
        score += 2.0
    return score / max(1.0, math.sqrt(len(tokens) + 1))


def chunk_text(text: str, chunk_size: int = 1200, overlap: int = 160) -> List[str]:
    clean = re.sub(r"\r\n?", "\n", str(text or "")).strip()
    clean = re.sub(r"\n{3,}", "\n\n", clean)
    if not clean:
        return []
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", clean) if p.strip()]
    chunks: List[str] = []
    current = ""
    for paragraph in paragraphs:
        if len(current) + len(paragraph) + 2 <= chunk_size:
            current = (current + "\n\n" + paragraph).strip()
            continue
        if current:
            chunks.append(current)
        if len(paragraph) <= chunk_size:
            current = paragraph
        else:
            start = 0
            while start < len(paragraph):
                piece = paragraph[start : start + chunk_size].strip()
                if piece:
                    chunks.append(piece)
                start += max(200, chunk_size - overlap)
            current = ""
    if current:
        chunks.append(current)
    return chunks[:3000]


# =========================
# Premium Knowledge Base helpers
# =========================

def _stable_hash(text: str) -> str:
    return hashlib.sha256(str(text or "").encode("utf-8", "ignore")).hexdigest()


def _normalize_collection(value: str) -> str:
    clean = re.sub(r"\s+", " ", str(value or "Default")).strip()
    return clean[:80] or "Default"


def _normalize_tags(value: str) -> str:
    raw = str(value or "")
    items = []
    for part in re.split(r"[,;#\n]", raw):
        tag = re.sub(r"[^a-zA-Z0-9_\-\s]", "", part).strip().lower()
        tag = re.sub(r"\s+", "-", tag)
        if tag and tag not in items:
            items.append(tag[:40])
    return ",".join(items[:20])


def _line_heading(line: str) -> str:
    raw = str(line or "").strip()
    if not raw:
        return ""
    if re.match(r"^#{1,6}\s+", raw):
        return re.sub(r"^#{1,6}\s+", "", raw).strip()[:160]
    if re.match(r"^(bab|chapter|section|seksi|bagian)\s+[0-9ivxlcdm\.\-]+", raw, flags=re.I):
        return raw[:160]
    if re.match(r"^\d+(\.\d+){0,4}\s+\S", raw) and len(raw) <= 160:
        return raw[:160]
    if raw.isupper() and 6 <= len(raw) <= 120:
        return raw[:160]
    return ""


def chunk_text_records(text: str, chunk_size: int = 1400, overlap: int = 180) -> List[Dict[str, Any]]:
    """Chunk documents with lightweight heading/page metadata for citation quality."""
    clean = re.sub(r"\r\n?", "\n", str(text or "")).strip()
    clean = re.sub(r"\n{4,}", "\n\n", clean)
    if not clean:
        return []
    records: List[Dict[str, Any]] = []
    heading = ""
    page_label = ""
    position = 0
    buffer = ""
    buffer_start = 0

    def flush() -> None:
        nonlocal buffer, buffer_start
        piece = buffer.strip()
        if piece:
            records.append({"content": piece, "heading": heading, "page_label": page_label, "char_start": buffer_start, "char_end": buffer_start + len(piece)})
        buffer = ""
        buffer_start = position

    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", clean) if p.strip()]
    for para in paragraphs:
        first_line = para.split("\n", 1)[0].strip()
        if re.match(r"^\[Halaman\s+[^\]]+\]", first_line, flags=re.I):
            page_label = first_line.strip("[]")[:80]
        new_heading = _line_heading(first_line)
        if new_heading:
            heading = new_heading
        para_len = len(para)
        if not buffer:
            buffer_start = position
        if len(buffer) + para_len + 2 <= chunk_size:
            buffer = (buffer + "\n\n" + para).strip()
        else:
            flush()
            if para_len <= chunk_size:
                buffer = para
                buffer_start = position
            else:
                start = 0
                while start < para_len:
                    piece = para[start : start + chunk_size].strip()
                    if piece:
                        records.append({"content": piece, "heading": heading, "page_label": page_label, "char_start": position + start, "char_end": position + start + len(piece)})
                    start += max(240, chunk_size - overlap)
                buffer = ""
                buffer_start = position + para_len
        position += para_len + 2
    flush()
    return records[:5000]


def _escape_fts_query(query: str) -> str:
    terms = _tokenize(query)[:10]
    return " OR ".join('"' + t.replace('"', '""') + '"' for t in terms) or str(query or "").replace('"', '""')


def _format_kb_citation(item: Dict[str, Any]) -> str:
    title = str(item.get("title") or "Dokumen").strip()
    chunk_index = item.get("chunk_index")
    page = str(item.get("page_label") or "").strip()
    heading = str(item.get("heading") or "").strip()
    parts = [title]
    if page:
        parts.append(page)
    if heading:
        parts.append(heading)
    if chunk_index is not None:
        parts.append(f"chunk {chunk_index}")
    return " · ".join(parts)[:260]



def _source_domain(source: str) -> str:
    """Extract a compact domain label from a URL/source string."""
    raw = str(source or "").strip()
    if not raw:
        return ""
    try:
        from urllib.parse import urlparse
        parsed = urlparse(raw if re.match(r"^https?://", raw, flags=re.I) else "https://" + raw)
        host = (parsed.netloc or parsed.path.split("/", 1)[0]).lower()
        return host.replace("www.", "")[:160]
    except Exception:
        return raw[:160]


def _estimate_source_quality(source: str = "", tags: str = "", metadata: Optional[Dict[str, Any]] = None) -> float:
    """Heuristic 0-100 source credibility score used for KB reranking.

    It is intentionally transparent and editable: admins can override it via
    source_quality on sources/docs. The score is not a truth guarantee; it only
    helps prefer official, journal, and primary sources when multiple chunks match.
    """
    metadata = metadata or {}
    explicit = metadata.get("source_quality") or metadata.get("quality_score")
    try:
        if explicit is not None and str(explicit).strip() != "":
            return max(0.0, min(100.0, float(explicit)))
    except Exception:
        pass

    hay = f"{source} {tags} {_safe_json(metadata)}".lower()
    score = 55.0
    official_markers = [".go.id", ".gov", "who.int", "fao.org", "woah.org", "nih.gov", "cdc.gov", "nasa.gov", "mit.edu", "kemkes.go.id", "pertanian.go.id"]
    journal_markers = ["scimagojr", "scopus", "sinta", "springer", "elsevier", "wiley", "tandfonline", "mdpi", "frontiersin", "nature.com", "science.org", "pubmed"]
    news_markers = ["kompas", "cnn", "reuters", "bbc", "tempo", "detik", "antaranews", "cnbc", "theverge", "techcrunch", "wired"]
    low_markers = ["blogspot", "wordpress", "medium.com", "facebook.com", "tiktok.com", "x.com", "twitter.com", "instagram.com"]
    if any(x in hay for x in official_markers):
        score = max(score, 92.0)
    if any(x in hay for x in journal_markers):
        score = max(score, 88.0)
    if any(x in hay for x in news_markers):
        score = max(score, 76.0)
    if any(x in hay for x in ["rss", "news", "berita", "trends", "trending"]):
        score = max(score, 68.0)
    if any(x in hay for x in low_markers):
        score = min(score, 45.0)
    return max(0.0, min(100.0, score))


def _recency_boost(ts: float) -> float:
    """Small freshness boost for recent KB chunks; capped so credibility still matters."""
    try:
        age_days = max(0.0, (_utc_ts() - float(ts or 0)) / 86400.0)
    except Exception:
        return 0.0
    if age_days <= 2:
        return 0.08
    if age_days <= 7:
        return 0.05
    if age_days <= 30:
        return 0.025
    return 0.0


# =========================
# Knowledge base file extraction helpers
# =========================

def _decode_bytes(data: bytes) -> str:
    raw = data or b""
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return raw.decode(enc, "ignore")
        except Exception:
            continue
    return raw.decode("utf-8", "ignore")


def _csv_bytes_to_text(data: bytes) -> str:
    text = _decode_bytes(data)
    reader = csv.reader(io.StringIO(text))
    lines: List[str] = []
    for idx, row in enumerate(reader):
        if idx >= 5000:
            lines.append("[CSV dipotong: lebih dari 5000 baris]")
            break
        clean = [str(cell or "").strip() for cell in row]
        if any(clean):
            lines.append(" | ".join(clean))
    return "\n".join(lines)


def _json_bytes_to_text(data: bytes) -> str:
    text = _decode_bytes(data)
    try:
        obj = json.loads(text)
        return json.dumps(obj, ensure_ascii=False, indent=2)
    except Exception:
        return text


def extract_text_from_file_bytes(filename: str, data: bytes) -> Tuple[str, str]:
    """Extract text from common knowledge-base uploads with optional dependencies.

    Supported without extra dependencies: txt, md, csv, json, log, py, js, html, css, xml.
    Supported if libraries exist in the app environment: pdf via pypdf/PyPDF2, docx via python-docx,
    xlsx via openpyxl. On failure, it returns a clear text note instead of crashing Streamlit.
    """
    name = str(filename or "uploaded_file").strip() or "uploaded_file"
    suffix = Path(name).suffix.lower()
    raw = data or b""
    if not raw:
        return "", "empty"

    text_suffixes = {".txt", ".md", ".markdown", ".log", ".py", ".js", ".ts", ".html", ".css", ".xml", ".jsonl"}
    if suffix in text_suffixes:
        return _decode_bytes(raw), "text"
    if suffix == ".csv":
        return _csv_bytes_to_text(raw), "csv"
    if suffix == ".json":
        return _json_bytes_to_text(raw), "json"

    if suffix == ".pdf":
        try:
            try:
                from pypdf import PdfReader  # type: ignore
            except Exception:
                from PyPDF2 import PdfReader  # type: ignore
            reader = PdfReader(io.BytesIO(raw))
            pages = []
            for i, page in enumerate(reader.pages[:250], start=1):
                try:
                    txt = page.extract_text() or ""
                except Exception:
                    txt = ""
                if txt.strip():
                    pages.append(f"[Halaman {i}]\n{txt.strip()}")
            return "\n\n".join(pages), "pdf"
        except Exception as exc:
            return f"[Gagal ekstrak PDF: {exc}]", "pdf_error"

    if suffix == ".docx":
        try:
            import docx  # type: ignore
            doc = docx.Document(io.BytesIO(raw))
            parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            return "\n".join(parts), "docx"
        except Exception as exc:
            return f"[Gagal ekstrak DOCX: {exc}]", "docx_error"

    if suffix in {".xlsx", ".xlsm"}:
        try:
            from openpyxl import load_workbook  # type: ignore
            wb = load_workbook(io.BytesIO(raw), data_only=True, read_only=True)
            lines: List[str] = []
            for ws in wb.worksheets[:20]:
                lines.append(f"[Sheet: {ws.title}]")
                for r_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
                    if r_idx > 5000:
                        lines.append("[Sheet dipotong: lebih dari 5000 baris]")
                        break
                    vals = [str(v).strip() if v is not None else "" for v in row]
                    if any(vals):
                        lines.append(" | ".join(vals))
            return "\n".join(lines), "xlsx"
        except Exception as exc:
            return f"[Gagal ekstrak XLSX: {exc}]", "xlsx_error"

    return _decode_bytes(raw), "binary_text_fallback"


INJECTION_PATTERNS = [
    r"(?i)abaikan\s+instruksi", r"(?i)ignore\s+(all\s+)?previous", r"(?i)developer\s+message",
    r"(?i)system\s+prompt", r"(?i)bocorkan\s+(secret|api|token|key)", r"(?i)reveal\s+(secret|api|token|key)",
    r"(?i)hapus\s+semua\s+aturan", r"(?i)gunakan\s+model\s+mahal\s+terus", r"(?i)jangan\s+ikuti\s+aturan",
]


def sanitize_non_instruction_context(text: str, limit: int = 4000) -> str:
    """Make memory/RAG context safer: it remains context, never instructions."""
    clean = str(text or "").replace("\x00", " ")
    for pattern in INJECTION_PATTERNS:
        clean = re.sub(pattern, "[fragmen instruksi tidak tepercaya dihapus]", clean)
    clean = re.sub(r"[ \t]+", " ", clean)
    clean = re.sub(r"\n{4,}", "\n\n", clean)
    return clean.strip()[: max(200, int(limit or 4000))]


def make_response_cache_key(
    *,
    model: str,
    system_prompt: str,
    user_text: str,
    memory_text: str,
    intent: str,
    route_signature: str = "",
) -> str:
    blob = json.dumps(
        {
            "model": model,
            "system": str(system_prompt or "")[:800],
            "user": str(user_text or ""),
            "memory_hash": hashlib.sha256(str(memory_text or "").encode("utf-8")).hexdigest(),
            "intent": intent,
            "route": route_signature,
        },
        ensure_ascii=False,
        sort_keys=True,
    )
    return hashlib.sha256(blob.encode("utf-8")).hexdigest()


def _tier_cost_penalty(model: str) -> float:
    tier = model_cost_tier(model)
    return {"cheap": 0.02, "medium": 0.12, "expensive": 0.38, "ultra": 0.95, "unknown": 0.28}.get(tier, 0.28)


def _base_model_bias(model: str, intent: str) -> float:
    m = str(model or "").lower()
    bias = 0.50
    if intent == "quick_chat":
        if any(x in m for x in ["flash", "nano", "mini", "haiku", "instant", "fast"]):
            bias += 0.16
    elif intent == "coding":
        if any(x in m for x in ["coder", "codex", "deepseek", "qwen"]):
            bias += 0.22
        if any(x in m for x in ["sonnet", "gpt-5.1", "gpt-5.2", "gpt-5.3", "gpt-5.4", "gpt-5.5"]):
            bias += 0.12
    elif intent in {"academic", "research", "deep_reasoning", "document_question"}:
        if any(x in m for x in ["sonnet", "qwen", "gpt-5.1", "gpt-5.2", "gpt-5.4", "gpt-5.5", "deepseek", "glm", "kimi"]):
            bias += 0.18
    elif intent == "creative":
        if any(x in m for x in ["claude", "gemini", "haiku", "sonnet", "gpt"]):
            bias += 0.14
    else:
        if any(x in m for x in ["mini", "flash", "nano", "haiku"]):
            bias += 0.08
    if ":free" in m or ":fast" in m:
        bias += 0.04
    if ":slow" in m:
        bias -= 0.03
    return max(0.05, min(0.95, bias))


def adaptive_token_budget_for_intent(intent: str, user_text: str, base: int = 1800) -> int:
    base = int(base or 1800)
    words = len(str(user_text or "").split())
    budgets = {
        "quick_chat": 800,
        "creative": 1400,
        "calculation": 1800,
        "coding": 3600,
        "academic": 4200,
        "document_question": 4800,
        "research": 4200,
        "critical_current": 4800,
        "health": 3600,
        "livestock": 3600,
        "deep_reasoning": 4200,
        "general": 2400,
    }
    target = budgets.get(intent, base)
    if words > 300:
        target = max(target, 4800)
    elif words > 140:
        target = max(target, 3600)
    return max(500, min(max(base, target), 6500))


def extract_quality_score(meta: Dict[str, Any], answer: str = "") -> float:
    try:
        if meta and meta.get("quality_score") is not None:
            return float(meta.get("quality_score") or 0)
    except Exception:
        pass
    words = len(str(answer or "").split())
    if words <= 0:
        return 0.0
    score = 0.42
    if words >= 20:
        score += 0.16
    if words >= 80:
        score += 0.10
    if "maaf" in str(answer or "").lower() and words < 50:
        score -= 0.08
    return max(0.0, min(1.0, score))


# =========================
# Store / schema
# =========================

@dataclass
class PowerStore:
    db_path: str = DEFAULT_POWER_DB

    def __post_init__(self) -> None:
        self.db_path = str(self.db_path or DEFAULT_POWER_DB)
        self._db_guard_recovered = False
        self._ensure_ready_before_open()
        try:
            self._ensure_schema()
        except sqlite3.DatabaseError as exc:
            # Jika SQLite mendeteksi "database disk image is malformed" atau error
            # sejenis, pulihkan otomatis dari backup valid terakhir lalu ulangi init schema.
            if self._looks_like_malformed_db_error(exc):
                self._restore_or_quarantine_database(reason=str(exc)[:300])
                self._ensure_schema()
            else:
                raise

    def _looks_like_malformed_db_error(self, exc: Exception) -> bool:
        message = str(exc or "").lower()
        return any(token in message for token in [
            "database disk image is malformed",
            "file is not a database",
            "database is locked",
            "disk i/o error",
            "malformed",
            "corrupt",
        ])

    def _ensure_ready_before_open(self) -> None:
        if ensure_database_ready is None:
            return
        try:
            ensure_database_ready(
                self.db_path,
                default_backup_dir(),
                auto_restore=True,
                create_periodic_backup=False,
                max_backups=default_max_backups(),
            )
        except Exception:
            # Jangan matikan aplikasi hanya karena guard gagal; _ensure_schema masih akan mencoba membuka DB.
            pass

    def _restore_or_quarantine_database(self, reason: str = "") -> None:
        if ensure_database_ready is None:
            raise sqlite3.DatabaseError(reason or "database error")
        result = ensure_database_ready(
            self.db_path,
            default_backup_dir(),
            auto_restore=True,
            create_periodic_backup=False,
            max_backups=default_max_backups(),
        )
        self._db_guard_recovered = True
        if not result.ok:
            raise sqlite3.DatabaseError(result.message or reason or "database restore failed")

    def _connect(self) -> sqlite3.Connection:
        path = Path(self.db_path)
        if path.parent and str(path.parent) not in {"", "."}:
            path.parent.mkdir(parents=True, exist_ok=True)
        conn = sqlite3.connect(self.db_path, timeout=30)
        conn.row_factory = sqlite3.Row
        return conn

    def _ensure_schema(self) -> None:
        with self._connect() as conn:
            conn.executescript(
                """
                PRAGMA journal_mode=WAL;
                CREATE TABLE IF NOT EXISTS memories (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    user_id TEXT NOT NULL DEFAULT 'global',
                    text TEXT NOT NULL,
                    tags TEXT DEFAULT '',
                    created_at REAL NOT NULL
                );
                CREATE INDEX IF NOT EXISTS idx_memories_user ON memories(user_id, created_at);

                CREATE TABLE IF NOT EXISTS documents (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    title TEXT NOT NULL,
                    source TEXT DEFAULT '',
                    created_at REAL NOT NULL
                );
                CREATE TABLE IF NOT EXISTS chunks (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    doc_id INTEGER NOT NULL,
                    chunk_index INTEGER NOT NULL,
                    content TEXT NOT NULL,
                    terms TEXT DEFAULT '',
                    created_at REAL NOT NULL,
                    FOREIGN KEY(doc_id) REFERENCES documents(id) ON DELETE CASCADE
                );
                CREATE INDEX IF NOT EXISTS idx_chunks_doc ON chunks(doc_id, chunk_index);

                CREATE TABLE IF NOT EXISTS interactions (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    ts REAL NOT NULL,
                    user_id TEXT DEFAULT 'public',
                    channel TEXT DEFAULT 'web',
                    intent TEXT DEFAULT '',
                    model TEXT DEFAULT '',
                    input_tokens INTEGER DEFAULT 0,
                    output_tokens INTEGER DEFAULT 0,
                    cost_idr REAL DEFAULT 0,
                    latency_seconds REAL DEFAULT 0,
                    success INTEGER DEFAULT 1,
                    question_preview TEXT DEFAULT '',
                    answer_preview TEXT DEFAULT '',
                    meta_json TEXT DEFAULT '{}'
                );
                CREATE INDEX IF NOT EXISTS idx_interactions_ts ON interactions(ts);
                CREATE INDEX IF NOT EXISTS idx_interactions_model ON interactions(model);
                CREATE INDEX IF NOT EXISTS idx_interactions_intent ON interactions(intent, ts);

                CREATE TABLE IF NOT EXISTS benchmarks (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    ts REAL NOT NULL,
                    model TEXT NOT NULL,
                    task TEXT NOT NULL,
                    score REAL DEFAULT 0,
                    latency_seconds REAL DEFAULT 0,
                    success INTEGER DEFAULT 0,
                    error TEXT DEFAULT '',
                    meta_json TEXT DEFAULT '{}'
                );
                CREATE INDEX IF NOT EXISTS idx_benchmarks_model ON benchmarks(model, ts);

                CREATE TABLE IF NOT EXISTS response_cache (
                    cache_key TEXT PRIMARY KEY,
                    ts REAL NOT NULL,
                    expires_at REAL NOT NULL,
                    model TEXT DEFAULT '',
                    intent TEXT DEFAULT '',
                    answer TEXT NOT NULL,
                    meta_json TEXT DEFAULT '{}'
                );
                CREATE INDEX IF NOT EXISTS idx_response_cache_expires ON response_cache(expires_at);

                CREATE TABLE IF NOT EXISTS semantic_response_cache (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    ts REAL NOT NULL,
                    expires_at REAL NOT NULL,
                    user_id TEXT DEFAULT 'public',
                    channel TEXT DEFAULT 'web',
                    intent TEXT DEFAULT '',
                    question TEXT NOT NULL,
                    question_terms TEXT DEFAULT '',
                    answer TEXT NOT NULL,
                    meta_json TEXT DEFAULT '{}'
                );
                CREATE INDEX IF NOT EXISTS idx_semantic_cache_intent ON semantic_response_cache(intent, expires_at);
                CREATE INDEX IF NOT EXISTS idx_semantic_cache_ts ON semantic_response_cache(ts);

                CREATE TABLE IF NOT EXISTS retrieval_eval_reports (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    ts REAL NOT NULL,
                    user_id TEXT DEFAULT 'public',
                    channel TEXT DEFAULT 'web',
                    intent TEXT DEFAULT '',
                    question TEXT DEFAULT '',
                    search_query TEXT DEFAULT '',
                    source_count INTEGER DEFAULT 0,
                    precision_estimate REAL DEFAULT 0,
                    avg_similarity REAL DEFAULT 0,
                    latency_seconds REAL DEFAULT 0,
                    meta_json TEXT DEFAULT '{}'
                );
                CREATE INDEX IF NOT EXISTS idx_retrieval_eval_ts ON retrieval_eval_reports(ts);
                CREATE INDEX IF NOT EXISTS idx_retrieval_eval_intent ON retrieval_eval_reports(intent, ts);

                CREATE TABLE IF NOT EXISTS source_runtime_stats (
                    source TEXT PRIMARY KEY,
                    total_requests INTEGER DEFAULT 0,
                    success_count INTEGER DEFAULT 0,
                    failure_count INTEGER DEFAULT 0,
                    total_latency REAL DEFAULT 0,
                    total_quality REAL DEFAULT 0,
                    last_error TEXT DEFAULT '',
                    last_seen_at REAL DEFAULT 0,
                    status TEXT DEFAULT 'active',
                    updated_at REAL NOT NULL
                );
                CREATE INDEX IF NOT EXISTS idx_source_runtime_status ON source_runtime_stats(status, updated_at);

                CREATE TABLE IF NOT EXISTS performance_events (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    ts REAL NOT NULL,
                    event_type TEXT DEFAULT '',
                    severity TEXT DEFAULT 'info',
                    title TEXT DEFAULT '',
                    detail TEXT DEFAULT '',
                    meta_json TEXT DEFAULT '{}'
                );
                CREATE INDEX IF NOT EXISTS idx_performance_events_ts ON performance_events(ts);

                CREATE TABLE IF NOT EXISTS model_scores (
                    model TEXT NOT NULL,
                    intent TEXT NOT NULL,
                    total_requests INTEGER DEFAULT 0,
                    success_count INTEGER DEFAULT 0,
                    error_count INTEGER DEFAULT 0,
                    total_latency REAL DEFAULT 0,
                    total_quality REAL DEFAULT 0,
                    total_cost REAL DEFAULT 0,
                    last_success_at REAL DEFAULT 0,
                    last_error_at REAL DEFAULT 0,
                    updated_at REAL NOT NULL,
                    PRIMARY KEY(model, intent)
                );
                CREATE INDEX IF NOT EXISTS idx_model_scores_intent ON model_scores(intent, updated_at);

                CREATE TABLE IF NOT EXISTS circuit_breakers (
                    model TEXT PRIMARY KEY,
                    failure_count INTEGER DEFAULT 0,
                    open_until REAL DEFAULT 0,
                    last_error TEXT DEFAULT '',
                    updated_at REAL NOT NULL
                );

                CREATE TABLE IF NOT EXISTS user_feedback (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    ts REAL NOT NULL,
                    interaction_id INTEGER DEFAULT 0,
                    user_id TEXT DEFAULT 'public',
                    rating INTEGER DEFAULT 0,
                    label TEXT DEFAULT '',
                    comment TEXT DEFAULT '',
                    created_by TEXT DEFAULT '',
                    meta_json TEXT DEFAULT '{}'
                );
                CREATE INDEX IF NOT EXISTS idx_user_feedback_ts ON user_feedback(ts);
                CREATE INDEX IF NOT EXISTS idx_user_feedback_interaction ON user_feedback(interaction_id);

                CREATE TABLE IF NOT EXISTS knowledge_gaps (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    ts REAL NOT NULL,
                    user_id TEXT DEFAULT 'public',
                    channel TEXT DEFAULT 'web',
                    intent TEXT DEFAULT '',
                    question TEXT NOT NULL,
                    reason TEXT DEFAULT '',
                    status TEXT DEFAULT 'open',
                    priority INTEGER DEFAULT 3,
                    suggested_query TEXT DEFAULT '',
                    meta_json TEXT DEFAULT '{}'
                );
                CREATE INDEX IF NOT EXISTS idx_knowledge_gaps_status ON knowledge_gaps(status, ts);
                CREATE INDEX IF NOT EXISTS idx_knowledge_gaps_intent ON knowledge_gaps(intent, ts);

                CREATE TABLE IF NOT EXISTS answer_templates (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    ts REAL NOT NULL,
                    title TEXT NOT NULL,
                    intent TEXT DEFAULT 'general',
                    trigger_query TEXT DEFAULT '',
                    answer TEXT NOT NULL,
                    tags TEXT DEFAULT '',
                    active INTEGER DEFAULT 1,
                    uses INTEGER DEFAULT 0,
                    updated_at REAL DEFAULT 0
                );
                CREATE INDEX IF NOT EXISTS idx_answer_templates_intent ON answer_templates(intent, active, updated_at);

                CREATE TABLE IF NOT EXISTS user_profiles (
                    user_id TEXT PRIMARY KEY,
                    answer_mode TEXT DEFAULT 'auto',
                    preferences_json TEXT DEFAULT '{}',
                    notes TEXT DEFAULT '',
                    updated_at REAL NOT NULL
                );

                CREATE TABLE IF NOT EXISTS answer_quality_reports (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    ts REAL NOT NULL,
                    interaction_id INTEGER DEFAULT 0,
                    user_id TEXT DEFAULT 'public',
                    channel TEXT DEFAULT 'web',
                    intent TEXT DEFAULT '',
                    answer_mode TEXT DEFAULT 'auto',
                    score REAL DEFAULT 0,
                    level TEXT DEFAULT '',
                    needs_verification INTEGER DEFAULT 0,
                    verifier_model TEXT DEFAULT '',
                    verified INTEGER DEFAULT 0,
                    reasons_json TEXT DEFAULT '[]',
                    metrics_json TEXT DEFAULT '{}',
                    question_preview TEXT DEFAULT '',
                    answer_preview TEXT DEFAULT ''
                );
                CREATE INDEX IF NOT EXISTS idx_answer_quality_ts ON answer_quality_reports(ts);
                CREATE INDEX IF NOT EXISTS idx_answer_quality_score ON answer_quality_reports(score, ts);
                CREATE INDEX IF NOT EXISTS idx_answer_quality_user ON answer_quality_reports(user_id, ts);

                CREATE TABLE IF NOT EXISTS weekly_evaluations (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    ts REAL NOT NULL,
                    days INTEGER DEFAULT 7,
                    report TEXT NOT NULL,
                    metrics_json TEXT DEFAULT '{}'
                );
                CREATE INDEX IF NOT EXISTS idx_weekly_evaluations_ts ON weekly_evaluations(ts);

                CREATE TABLE IF NOT EXISTS exported_snapshots (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    ts REAL NOT NULL,
                    snapshot_type TEXT NOT NULL,
                    filename TEXT DEFAULT '',
                    rows_count INTEGER DEFAULT 0,
                    meta_json TEXT DEFAULT '{}'
                );

                CREATE TABLE IF NOT EXISTS source_quality_rules (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    pattern TEXT NOT NULL UNIQUE,
                    quality REAL NOT NULL DEFAULT 60,
                    note TEXT DEFAULT '',
                    updated_at REAL NOT NULL
                );

                CREATE TABLE IF NOT EXISTS current_claims (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    doc_id INTEGER DEFAULT 0,
                    ts REAL NOT NULL,
                    published_at REAL DEFAULT 0,
                    claim TEXT NOT NULL,
                    title TEXT DEFAULT '',
                    source TEXT DEFAULT '',
                    source_name TEXT DEFAULT '',
                    source_quality REAL DEFAULT 55,
                    freshness_score REAL DEFAULT 45,
                    category TEXT DEFAULT '',
                    keywords TEXT DEFAULT '',
                    metadata_json TEXT DEFAULT '{}',
                    FOREIGN KEY(doc_id) REFERENCES documents(id) ON DELETE CASCADE
                );
                CREATE INDEX IF NOT EXISTS idx_current_claims_ts ON current_claims(ts);
                CREATE INDEX IF NOT EXISTS idx_current_claims_doc ON current_claims(doc_id);
                CREATE INDEX IF NOT EXISTS idx_current_claims_quality ON current_claims(source_quality, freshness_score);

                CREATE TABLE IF NOT EXISTS issue_watchlist (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    topic TEXT NOT NULL UNIQUE,
                    category TEXT DEFAULT 'general',
                    priority INTEGER DEFAULT 3,
                    active INTEGER DEFAULT 1,
                    created_at REAL NOT NULL,
                    updated_at REAL NOT NULL,
                    meta_json TEXT DEFAULT '{}'
                );
                CREATE INDEX IF NOT EXISTS idx_issue_watchlist_active ON issue_watchlist(active, priority, updated_at);

                CREATE TABLE IF NOT EXISTS issue_events (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    ts REAL NOT NULL,
                    topic TEXT NOT NULL,
                    title TEXT DEFAULT '',
                    source TEXT DEFAULT '',
                    event_text TEXT NOT NULL,
                    event_time REAL DEFAULT 0,
                    source_quality REAL DEFAULT 55,
                    freshness_score REAL DEFAULT 45,
                    meta_json TEXT DEFAULT '{}'
                );
                CREATE INDEX IF NOT EXISTS idx_issue_events_topic ON issue_events(topic, ts);
                """
            )
            # Optional SQLite FTS5. Streamlit Cloud usually supports it, but keep fallback.
            try:
                conn.execute(
                    "CREATE VIRTUAL TABLE IF NOT EXISTS chunks_fts USING fts5(content, title, source, collection, tags, heading, page_label, content='')"
                )
            except Exception:
                pass

            for ddl in [
                "ALTER TABLE documents ADD COLUMN collection TEXT DEFAULT 'Default'",
                "ALTER TABLE documents ADD COLUMN tags TEXT DEFAULT ''",
                "ALTER TABLE documents ADD COLUMN doc_hash TEXT DEFAULT ''",
                "ALTER TABLE documents ADD COLUMN metadata_json TEXT DEFAULT '{}'",
                "ALTER TABLE documents ADD COLUMN pinned INTEGER DEFAULT 0",
                "ALTER TABLE documents ADD COLUMN updated_at REAL DEFAULT 0",
                "ALTER TABLE chunks ADD COLUMN heading TEXT DEFAULT ''",
                "ALTER TABLE chunks ADD COLUMN page_label TEXT DEFAULT ''",
                "ALTER TABLE chunks ADD COLUMN char_start INTEGER DEFAULT 0",
                "ALTER TABLE chunks ADD COLUMN char_end INTEGER DEFAULT 0",
                "ALTER TABLE documents ADD COLUMN source_quality REAL DEFAULT 55",
                "ALTER TABLE documents ADD COLUMN source_domain TEXT DEFAULT ''",
                "ALTER TABLE documents ADD COLUMN summary TEXT DEFAULT ''",
                "ALTER TABLE documents ADD COLUMN published_at REAL DEFAULT 0",
                "ALTER TABLE documents ADD COLUMN freshness_score REAL DEFAULT 45",
                "ALTER TABLE documents ADD COLUMN criticality_score REAL DEFAULT 0",
            ]:
                try:
                    conn.execute(ddl)
                except Exception:
                    pass
            try:
                conn.execute("CREATE INDEX IF NOT EXISTS idx_documents_collection ON documents(collection, created_at)")
                conn.execute("CREATE INDEX IF NOT EXISTS idx_documents_hash ON documents(doc_hash)")
                conn.execute("CREATE INDEX IF NOT EXISTS idx_documents_pinned ON documents(pinned, created_at)")
                conn.execute("CREATE INDEX IF NOT EXISTS idx_documents_quality ON documents(source_quality, created_at)")
                conn.execute("CREATE INDEX IF NOT EXISTS idx_documents_domain ON documents(source_domain)")
                conn.execute("CREATE INDEX IF NOT EXISTS idx_documents_freshness ON documents(freshness_score, published_at)")
                conn.execute("CREATE INDEX IF NOT EXISTS idx_documents_criticality ON documents(criticality_score, created_at)")
            except Exception:
                pass

    # Persistent memory
    def add_memory(self, text: str, user_id: str = "global", tags: str = "") -> int:
        clean = sanitize_non_instruction_context(str(text or "").strip(), limit=4000)
        if not clean:
            return 0
        with self._connect() as conn:
            cur = conn.execute(
                "INSERT INTO memories(user_id, text, tags, created_at) VALUES (?, ?, ?, ?)",
                (str(user_id or "global"), clean, str(tags or "")[:300], _utc_ts()),
            )
            return int(cur.lastrowid)

    def delete_memories_containing(self, keyword: str, user_id: Optional[str] = None) -> int:
        key = str(keyword or "").strip().lower()
        if not key:
            return 0
        with self._connect() as conn:
            if user_id:
                cur = conn.execute(
                    "DELETE FROM memories WHERE user_id = ? AND lower(text) LIKE ?",
                    (str(user_id), f"%{key}%"),
                )
            else:
                cur = conn.execute("DELETE FROM memories WHERE lower(text) LIKE ?", (f"%{key}%",))
            return int(cur.rowcount or 0)

    def search_memories(self, query: str, user_id: str = "global", limit: int = 8) -> List[Dict[str, Any]]:
        with self._connect() as conn:
            rows = conn.execute(
                "SELECT * FROM memories WHERE user_id IN (?, 'global') ORDER BY created_at DESC LIMIT 500",
                (str(user_id or "global"),),
            ).fetchall()
        scored: List[Tuple[float, Dict[str, Any]]] = []
        for row in rows:
            item = dict(row)
            score = _score_text(query, item.get("text", ""))
            if score > 0 or len(str(item.get("text", ""))) <= 220:
                scored.append((score, item))
        scored.sort(key=lambda x: (x[0], x[1].get("created_at", 0)), reverse=True)
        return [item for _, item in scored[: max(1, int(limit or 8))]]

    # Knowledge base / RAG
    def add_document(
        self,
        title: str,
        text: str,
        source: str = "manual",
        collection: str = "Default",
        tags: str = "",
        metadata: Optional[Dict[str, Any]] = None,
        replace_existing: bool = False,
        pinned: bool = False,
        source_quality: Optional[float] = None,
        summary: str = "",
    ) -> Tuple[int, int]:
        """Add a document with workspace/collection, tags, hashing, and citation-ready chunks."""
        title = str(title or "Dokumen tanpa judul").strip()[:240]
        source = str(source or "manual").strip()[:500]
        collection = _normalize_collection(collection)
        tags = _normalize_tags(tags)
        raw_text = str(text or "").strip()
        records = chunk_text_records(raw_text)
        if not records:
            return 0, 0
        doc_hash = _stable_hash(title + "\n" + source + "\n" + raw_text[:2_000_000])
        metadata = metadata or {}
        if source_quality is not None:
            metadata["source_quality"] = source_quality
        summary = sanitize_non_instruction_context(str(summary or metadata.get("summary") or ""), limit=1200)
        source_domain = _source_domain(source)
        computed_quality = _estimate_source_quality(source=source, tags=tags, metadata=metadata)
        freshness_info = calculate_freshness_score(metadata.get("published") or metadata.get("published_at") or "", metadata.get("scraped_at") or metadata.get("created_at") or "")
        try:
            published_ts = float(metadata.get("published_ts") or metadata.get("published_at_ts") or freshness_info.get("published_ts") or 0)
        except Exception:
            published_ts = 0.0
        try:
            freshness_score = float(metadata.get("freshness_score") if metadata.get("freshness_score") is not None else freshness_info.get("score", 45))
        except Exception:
            freshness_score = 45.0
        critical_detection = detect_critical_question(f"{title} {collection} {tags} {summary} {raw_text[:1200]}")
        try:
            criticality_score = float(metadata.get("criticality_score") if metadata.get("criticality_score") is not None else critical_detection.get("score", 0))
        except Exception:
            criticality_score = 0.0
        metadata.setdefault("freshness", freshness_info)
        metadata.setdefault("critical_detection", critical_detection)
        meta_json = _safe_json(metadata or {})
        now = _utc_ts()
        with self._connect() as conn:
            existing = conn.execute("SELECT id FROM documents WHERE doc_hash = ? LIMIT 1", (doc_hash,)).fetchone()
            if existing and not replace_existing:
                return int(existing["id"]), 0
            if existing and replace_existing:
                try:
                    self.delete_document(int(existing["id"]))
                except Exception:
                    pass
            cur = conn.execute(
                """
                INSERT INTO documents(title, source, collection, tags, doc_hash, metadata_json, pinned, created_at, updated_at, source_quality, source_domain, summary, published_at, freshness_score, criticality_score)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (title, source, collection, tags, doc_hash, meta_json, 1 if pinned else 0, now, now, computed_quality, source_domain, summary, published_ts, freshness_score, criticality_score),
            )
            doc_id = int(cur.lastrowid)
            inserted = 0
            for idx, record in enumerate(records):
                safe_chunk = sanitize_non_instruction_context(record.get("content", ""), limit=6500)
                if not safe_chunk:
                    continue
                terms = " ".join(sorted(set(_tokenize(safe_chunk + " " + title + " " + tags + " " + collection)))[:700])
                cur_chunk = conn.execute(
                    """
                    INSERT INTO chunks(doc_id, chunk_index, content, terms, heading, page_label, char_start, char_end, created_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (doc_id, idx, safe_chunk, terms, str(record.get("heading") or "")[:180], str(record.get("page_label") or "")[:80], int(record.get("char_start") or 0), int(record.get("char_end") or 0), now),
                )
                inserted += 1
                try:
                    conn.execute(
                        "INSERT INTO chunks_fts(rowid, content, title, source, collection, tags, heading, page_label) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                        (int(cur_chunk.lastrowid), safe_chunk, title, source, collection, tags, str(record.get("heading") or "")[:180], str(record.get("page_label") or "")[:80]),
                    )
                except Exception:
                    pass
            try:
                claims = metadata.get("claims") or []
                if isinstance(claims, list):
                    for claim_item in claims[:30]:
                        if isinstance(claim_item, dict):
                            claim_text = sanitize_non_instruction_context(str(claim_item.get("claim") or ""), limit=900)
                            keywords = ",".join(str(x)[:40] for x in (claim_item.get("keywords") or [])[:12]) if isinstance(claim_item.get("keywords"), list) else str(claim_item.get("keywords") or "")[:300]
                            source_name = str(claim_item.get("source_name") or metadata.get("source_name") or "")[:180]
                        else:
                            claim_text = sanitize_non_instruction_context(str(claim_item or ""), limit=900)
                            keywords = ""
                            source_name = str(metadata.get("source_name") or "")[:180]
                        if not claim_text:
                            continue
                        conn.execute(
                            """
                            INSERT INTO current_claims(doc_id, ts, published_at, claim, title, source, source_name, source_quality, freshness_score, category, keywords, metadata_json)
                            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                            """,
                            (doc_id, now, published_ts, claim_text, title, source, source_name, computed_quality, freshness_score, collection, keywords, _safe_json(claim_item if isinstance(claim_item, dict) else {})),
                        )
            except Exception:
                pass
            return doc_id, inserted

    def search_documents(self, query: str, limit: int = 5, collection: str = "", min_score: float = 0.0, include_pinned: bool = True) -> List[Dict[str, Any]]:
        """Hybrid KB search: FTS5 + lexical reranking + metadata weighting."""
        q = str(query or "").strip()
        if not q:
            return []
        limit = max(1, int(limit or 5))
        collection_filter = _normalize_collection(collection) if str(collection or "").strip() else ""
        results: Dict[int, Dict[str, Any]] = {}
        terms = _escape_fts_query(q)
        try:
            with self._connect() as conn:
                if collection_filter:
                    rows = conn.execute(
                        """
                        SELECT c.id, c.doc_id, c.chunk_index, c.content, c.heading, c.page_label,
                               d.title, d.source, d.collection, d.tags, d.pinned, d.source_quality, d.source_domain, d.freshness_score, d.published_at, d.criticality_score, c.created_at,
                               bm25(chunks_fts) AS fts_score
                        FROM chunks_fts
                        JOIN chunks c ON c.id = chunks_fts.rowid
                        JOIN documents d ON d.id = c.doc_id
                        WHERE chunks_fts MATCH ? AND d.collection = ?
                        ORDER BY fts_score ASC LIMIT ?
                        """,
                        (terms, collection_filter, limit * 5),
                    ).fetchall()
                else:
                    rows = conn.execute(
                        """
                        SELECT c.id, c.doc_id, c.chunk_index, c.content, c.heading, c.page_label,
                               d.title, d.source, d.collection, d.tags, d.pinned, d.source_quality, d.source_domain, d.freshness_score, d.published_at, d.criticality_score, c.created_at,
                               bm25(chunks_fts) AS fts_score
                        FROM chunks_fts
                        JOIN chunks c ON c.id = chunks_fts.rowid
                        JOIN documents d ON d.id = c.doc_id
                        WHERE chunks_fts MATCH ?
                        ORDER BY fts_score ASC LIMIT ?
                        """,
                        (terms, limit * 5),
                    ).fetchall()
            for row in rows:
                item = dict(row)
                lexical = _score_text(q, f"{item.get('title','')} {item.get('heading','')} {item.get('content','')} {item.get('tags','')}")
                fts_component = 1.0 / (1.0 + abs(float(item.get("fts_score") or 0)))
                score = (fts_component * 1.35) + lexical
                if item.get("pinned"):
                    score += 0.08
                try:
                    score += (float(item.get("source_quality") or 55) / 100.0) * 0.12
                except Exception:
                    pass
                score += _recency_boost(float(item.get("created_at") or 0))
                try:
                    score += (float(item.get("freshness_score") or 45) / 100.0) * 0.10
                    score += min(0.08, (float(item.get("criticality_score") or 0) / 100.0) * 0.08)
                except Exception:
                    pass
                item["score"] = round(score, 4)
                item["citation"] = _format_kb_citation(item)
                item["source_type"] = "fts5"
                if score >= min_score:
                    results[int(item["id"])] = item
        except Exception:
            pass

        with self._connect() as conn:
            if collection_filter:
                rows = conn.execute(
                    """
                    SELECT c.id, c.doc_id, c.chunk_index, c.content, c.heading, c.page_label,
                           d.title, d.source, d.collection, d.tags, d.pinned, d.source_quality, d.source_domain, d.freshness_score, d.published_at, d.criticality_score, c.created_at
                    FROM chunks c JOIN documents d ON d.id = c.doc_id
                    WHERE d.collection = ?
                    ORDER BY d.pinned DESC, c.created_at DESC LIMIT 2500
                    """,
                    (collection_filter,),
                ).fetchall()
            else:
                rows = conn.execute(
                    """
                    SELECT c.id, c.doc_id, c.chunk_index, c.content, c.heading, c.page_label,
                           d.title, d.source, d.collection, d.tags, d.pinned, d.source_quality, d.source_domain, d.freshness_score, d.published_at, d.criticality_score, c.created_at
                    FROM chunks c JOIN documents d ON d.id = c.doc_id
                    ORDER BY d.pinned DESC, c.created_at DESC LIMIT 2500
                    """
                ).fetchall()
        scored: List[Tuple[float, Dict[str, Any]]] = []
        for row in rows:
            item = dict(row)
            weighted_text = f"{item.get('title','')} {item.get('title','')} {item.get('heading','')} {item.get('tags','')} {item.get('collection','')} {item.get('content','')}"
            score = _score_text(q, weighted_text)
            if item.get("pinned"):
                score += 0.08
            try:
                score += (float(item.get("source_quality") or 55) / 100.0) * 0.12
            except Exception:
                pass
            score += _recency_boost(float(item.get("created_at") or 0))
            try:
                score += (float(item.get("freshness_score") or 45) / 100.0) * 0.10
                score += min(0.08, (float(item.get("criticality_score") or 0) / 100.0) * 0.08)
            except Exception:
                pass
            if score > min_score:
                item["score"] = round(score, 4)
                item["citation"] = _format_kb_citation(item)
                item["source_type"] = "lexical"
                scored.append((score, item))
        scored.sort(key=lambda x: x[0], reverse=True)
        for score, item in scored[: limit * 5]:
            existing = results.get(int(item["id"]))
            if not existing or float(item.get("score") or 0) > float(existing.get("score") or 0):
                results[int(item["id"])] = item
        final = list(results.values())
        final.sort(key=lambda x: (float(x.get("score") or 0), int(x.get("pinned") or 0)), reverse=True)
        return final[:limit]

    # Critical current knowledge / issue intelligence
    def search_current_claims(self, query: str, limit: int = 10, min_score: float = 0.0, days: int = 90) -> List[Dict[str, Any]]:
        q = str(query or "").strip()
        if not q:
            return []
        since = _utc_ts() - max(1, int(days or 90)) * 86400
        with self._connect() as conn:
            rows = conn.execute(
                """
                SELECT * FROM current_claims
                WHERE ts >= ? OR published_at >= ?
                ORDER BY freshness_score DESC, source_quality DESC, ts DESC
                LIMIT 2000
                """,
                (since, since),
            ).fetchall()
        scored: List[Tuple[float, Dict[str, Any]]] = []
        for row in rows:
            item = _row_to_dict(row)
            text = f"{item.get('claim','')} {item.get('title','')} {item.get('source_name','')} {item.get('keywords','')} {item.get('category','')}"
            score = _score_text(q, text)
            try:
                score += (float(item.get("source_quality") or 55) / 100.0) * 0.16
                score += (float(item.get("freshness_score") or 45) / 100.0) * 0.18
            except Exception:
                pass
            if score > min_score:
                item["score"] = round(score, 4)
                item["ts_wib"] = _timestamp_to_wib(float(item.get("ts") or 0))
                item["published_wib"] = _timestamp_to_wib(float(item.get("published_at") or 0))
                scored.append((score, item))
        scored.sort(key=lambda x: x[0], reverse=True)
        return [item for _, item in scored[:max(1, int(limit or 10))]]

    def build_current_issue_report(self, query: str, limit: int = 8) -> str:
        claims = self.search_current_claims(query, limit=limit, days=180)
        docs = self.search_documents(query, limit=5)
        return format_issue_report(query, claims=claims, docs=docs)

    def add_watch_topic(self, topic: str, category: str = "general", priority: int = 3, meta: Optional[Dict[str, Any]] = None) -> int:
        clean = sanitize_non_instruction_context(topic, limit=240)
        if not clean:
            return 0
        now = _utc_ts()
        with self._connect() as conn:
            cur = conn.execute(
                """
                INSERT INTO issue_watchlist(topic, category, priority, active, created_at, updated_at, meta_json)
                VALUES (?, ?, ?, 1, ?, ?, ?)
                ON CONFLICT(topic) DO UPDATE SET category=excluded.category, priority=excluded.priority, active=1, updated_at=excluded.updated_at, meta_json=excluded.meta_json
                """,
                (clean, str(category or "general")[:80], max(1, min(5, int(priority or 3))), now, now, _safe_json(meta or {})),
            )
            row = conn.execute("SELECT id FROM issue_watchlist WHERE topic = ?", (clean,)).fetchone()
            return int(_row_value(row, "id", 0, cur.lastrowid or 0) or 0)

    def list_watch_topics(self, active_only: bool = True, limit: int = 50) -> List[Dict[str, Any]]:
        with self._connect() as conn:
            if active_only:
                rows = conn.execute("SELECT * FROM issue_watchlist WHERE active = 1 ORDER BY priority ASC, updated_at DESC LIMIT ?", (max(1, int(limit or 50)),)).fetchall()
            else:
                rows = conn.execute("SELECT * FROM issue_watchlist ORDER BY active DESC, priority ASC, updated_at DESC LIMIT ?", (max(1, int(limit or 50)),)).fetchall()
        out = []
        for row in rows:
            item = _row_to_dict(row)
            item["updated_at_wib"] = _timestamp_to_wib(float(item.get("updated_at") or 0))
            out.append(item)
        return out

    def remove_watch_topic(self, watch_id: int) -> bool:
        with self._connect() as conn:
            cur = conn.execute("UPDATE issue_watchlist SET active = 0, updated_at = ? WHERE id = ?", (_utc_ts(), int(watch_id or 0)))
            return bool(cur.rowcount)

    def daily_current_briefing(self, days: int = 1, limit: int = 12) -> str:
        since = _utc_ts() - max(1, int(days or 1)) * 86400
        with self._connect() as conn:
            rows = conn.execute(
                """
                SELECT * FROM current_claims
                WHERE ts >= ? OR published_at >= ?
                ORDER BY source_quality DESC, freshness_score DESC, ts DESC
                LIMIT ?
                """,
                (since, since, max(1, int(limit or 12))),
            ).fetchall()
        claims = [_row_to_dict(row) for row in rows]
        lines = [
            "📌 Daily Intelligence Briefing Adioranye",
            f"Periode: {days} hari terakhir",
            f"Tanggal cek: {now_wib_text()}",
            "",
        ]
        if not claims:
            lines.append("Belum ada klaim/fakta baru yang terekam di current_claims. Jalankan /update atau cek sumber KB.")
            return "\n".join(lines)
        lines.append("Update penting dari KB:")
        for idx, item in enumerate(claims, start=1):
            claim = re.sub(r"\s+", " ", str(item.get("claim") or "")).strip()[:360]
            title = str(item.get("title") or item.get("source_name") or "-")[:150]
            lines.append(
                f"{idx}. {claim}\n"
                f"   Sumber: {title} | kualitas {item.get('source_quality')} | freshness {item.get('freshness_score')} | { _timestamp_to_wib(float(item.get('published_at') or item.get('ts') or 0)) }"
            )
        lines.append("\nGunakan /cek isu <topik> untuk melihat bukti dan dokumen pendukung pada satu isu tertentu.")
        return "\n".join(lines)

    # Learning loop / feedback / knowledge gaps
    def record_feedback(
        self,
        interaction_id: int = 0,
        rating: int = 0,
        label: str = "",
        comment: str = "",
        user_id: str = "public",
        created_by: str = "user",
        meta: Optional[Dict[str, Any]] = None,
    ) -> int:
        """Store user/admin feedback so routing and answer templates can improve."""
        rating_value = max(-1, min(1, int(rating or 0)))
        now = _utc_ts()
        with self._connect() as conn:
            cur = conn.execute(
                """
                INSERT INTO user_feedback(ts,interaction_id,user_id,rating,label,comment,created_by,meta_json)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (now, int(interaction_id or 0), str(user_id or "public")[:120], rating_value, str(label or "")[:80], str(comment or "")[:1200], str(created_by or "user")[:80], _safe_json(meta or {})[:4000]),
            )
            # Nudge the historical model quality score when we know which interaction is being rated.
            row = conn.execute("SELECT model, intent, latency_seconds, cost_idr FROM interactions WHERE id = ?", (int(interaction_id or 0),)).fetchone()
        if row is not None and rating_value != 0:
            quality = 0.85 if rating_value > 0 else 0.15
            try:
                self.update_model_score(
                    model=str(_row_value(row, "model", 0, "")),
                    intent=str(_row_value(row, "intent", 1, "general")),
                    success=rating_value > 0,
                    latency_seconds=float(_row_value(row, "latency_seconds", 2, 0) or 0),
                    quality_score=quality,
                    cost_idr=float(_row_value(row, "cost_idr", 3, 0) or 0),
                )
            except Exception:
                pass
        return int(cur.lastrowid)

    def feedback_summary(self, days: int = 30) -> Dict[str, Any]:
        since = _utc_ts() - max(1, int(days or 30)) * 86400
        with self._connect() as conn:
            total = conn.execute(
                "SELECT COUNT(*) AS n, COALESCE(SUM(CASE WHEN rating > 0 THEN 1 ELSE 0 END),0) AS up, COALESCE(SUM(CASE WHEN rating < 0 THEN 1 ELSE 0 END),0) AS down FROM user_feedback WHERE ts >= ?",
                (since,),
            ).fetchone()
            by_label = conn.execute(
                "SELECT label, COUNT(*) AS jumlah, COALESCE(AVG(rating),0) AS avg_rating FROM user_feedback WHERE ts >= ? GROUP BY label ORDER BY jumlah DESC LIMIT 20",
                (since,),
            ).fetchall()
        return {
            "days": days,
            "total": int(_row_value(total, "n", 0, 0) or 0),
            "positive": int(_row_value(total, "up", 1, 0) or 0),
            "negative": int(_row_value(total, "down", 2, 0) or 0),
            "by_label": [_row_to_dict(row) for row in by_label],
        }

    def recent_interactions(self, limit: int = 50, only_negative: bool = False) -> List[Dict[str, Any]]:
        where = ""
        params: Tuple[Any, ...] = (max(1, int(limit or 50)),)
        if only_negative:
            where = "WHERE EXISTS (SELECT 1 FROM user_feedback f WHERE f.interaction_id = interactions.id AND f.rating < 0)"
        with self._connect() as conn:
            rows = conn.execute(
                f"""
                SELECT id, ts, user_id, channel, intent, model, success, question_preview, answer_preview, latency_seconds, cost_idr
                FROM interactions {where}
                ORDER BY ts DESC LIMIT ?
                """,
                params,
            ).fetchall()
        out = []
        for row in rows:
            item = _row_to_dict(row)
            item["ts_wib"] = _timestamp_to_wib(float(item.get("ts") or 0))
            out.append(item)
        return out

    def log_knowledge_gap(
        self,
        question: str,
        reason: str = "low_kb_coverage",
        intent: str = "general",
        user_id: str = "public",
        channel: str = "web",
        priority: int = 3,
        suggested_query: str = "",
        meta: Optional[Dict[str, Any]] = None,
    ) -> int:
        clean_q = sanitize_non_instruction_context(question, limit=1200)
        if not clean_q:
            return 0
        with self._connect() as conn:
            cur = conn.execute(
                """
                INSERT INTO knowledge_gaps(ts,user_id,channel,intent,question,reason,status,priority,suggested_query,meta_json)
                VALUES (?, ?, ?, ?, ?, ?, 'open', ?, ?, ?)
                """,
                (_utc_ts(), str(user_id or "public")[:120], str(channel or "web")[:40], str(intent or "general")[:80], clean_q, str(reason or "")[:120], max(1, min(5, int(priority or 3))), str(suggested_query or clean_q)[:500], _safe_json(meta or {})[:4000]),
            )
            return int(cur.lastrowid)

    def list_knowledge_gaps(self, status: str = "open", limit: int = 50) -> List[Dict[str, Any]]:
        params: Tuple[Any, ...]
        where = ""
        if status:
            where = "WHERE status = ?"
            params = (str(status)[:40], max(1, int(limit or 50)))
        else:
            params = (max(1, int(limit or 50)),)
        with self._connect() as conn:
            rows = conn.execute(
                f"SELECT * FROM knowledge_gaps {where} ORDER BY priority ASC, ts DESC LIMIT ?",
                params,
            ).fetchall()
        out = []
        for row in rows:
            item = _row_to_dict(row)
            item["ts_wib"] = _timestamp_to_wib(float(item.get("ts") or 0))
            out.append(item)
        return out

    def update_knowledge_gap_status(self, gap_id: int, status: str = "done") -> bool:
        with self._connect() as conn:
            cur = conn.execute("UPDATE knowledge_gaps SET status = ? WHERE id = ?", (str(status or "done")[:40], int(gap_id or 0)))
            return bool(cur.rowcount)

    def save_answer_template(self, title: str, trigger_query: str, answer: str, intent: str = "general", tags: str = "") -> int:
        body = sanitize_non_instruction_context(answer, limit=8000)
        if not body:
            return 0
        now = _utc_ts()
        with self._connect() as conn:
            cur = conn.execute(
                """
                INSERT INTO answer_templates(ts,title,intent,trigger_query,answer,tags,active,uses,updated_at)
                VALUES (?, ?, ?, ?, ?, ?, 1, 0, ?)
                """,
                (now, str(title or "Template jawaban")[:200], str(intent or "general")[:80], str(trigger_query or "")[:500], body, _normalize_tags(tags), now),
            )
            return int(cur.lastrowid)

    def search_answer_templates(self, query: str, intent: str = "", limit: int = 3) -> List[Dict[str, Any]]:
        with self._connect() as conn:
            if intent:
                rows = conn.execute("SELECT * FROM answer_templates WHERE active = 1 AND intent IN (?, 'general') ORDER BY updated_at DESC LIMIT 500", (str(intent)[:80],)).fetchall()
            else:
                rows = conn.execute("SELECT * FROM answer_templates WHERE active = 1 ORDER BY updated_at DESC LIMIT 500").fetchall()
        scored: List[Tuple[float, Dict[str, Any]]] = []
        for row in rows:
            item = _row_to_dict(row)
            score = _score_text(query, f"{item.get('title','')} {item.get('trigger_query','')} {item.get('tags','')}")
            if score > 0:
                scored.append((score, item))
        scored.sort(key=lambda x: x[0], reverse=True)
        return [item for _, item in scored[:max(1, int(limit or 3))]]

    def learning_dashboard(self, days: int = 14) -> Dict[str, Any]:
        since = _utc_ts() - max(1, int(days or 14)) * 86400
        with self._connect() as conn:
            intents = conn.execute("SELECT intent, COUNT(*) AS jumlah, COALESCE(AVG(success),0) AS success_rate FROM interactions WHERE ts >= ? GROUP BY intent ORDER BY jumlah DESC LIMIT 20", (since,)).fetchall()
            repeated = conn.execute("SELECT lower(question_preview) AS pertanyaan, COUNT(*) AS jumlah FROM interactions WHERE ts >= ? AND length(question_preview) > 12 GROUP BY lower(question_preview) HAVING jumlah > 1 ORDER BY jumlah DESC LIMIT 20", (since,)).fetchall()
            gaps = conn.execute("SELECT status, COUNT(*) AS jumlah FROM knowledge_gaps WHERE ts >= ? GROUP BY status", (since,)).fetchall()
        return {
            "days": days,
            "intents": [_row_to_dict(row) for row in intents],
            "repeated_questions": [_row_to_dict(row) for row in repeated],
            "gap_status": [_row_to_dict(row) for row in gaps],
            "feedback": self.feedback_summary(days=days),
        }

    # Quality control / user profile / export helpers
    def get_user_profile(self, user_id: str = "public") -> Dict[str, Any]:
        uid = str(user_id or "public")[:120]
        with self._connect() as conn:
            row = conn.execute("SELECT * FROM user_profiles WHERE user_id = ?", (uid,)).fetchone()
        if not row:
            return {"user_id": uid, "answer_mode": "auto", "preferences": {}, "notes": ""}
        item = _row_to_dict(row)
        item["preferences"] = _safe_json_loads(item.get("preferences_json"))
        return item

    def set_user_answer_mode(self, user_id: str = "public", answer_mode: str = "auto") -> str:
        uid = str(user_id or "public")[:120]
        mode = normalize_answer_mode(answer_mode)
        now = _utc_ts()
        with self._connect() as conn:
            conn.execute(
                """
                INSERT INTO user_profiles(user_id,answer_mode,preferences_json,notes,updated_at)
                VALUES (?, ?, '{}', '', ?)
                ON CONFLICT(user_id) DO UPDATE SET answer_mode=excluded.answer_mode, updated_at=excluded.updated_at
                """,
                (uid, mode, now),
            )
        return mode

    def get_user_answer_mode(self, user_id: str = "public", default: str = "auto") -> str:
        try:
            profile = self.get_user_profile(user_id)
            return normalize_answer_mode(profile.get("answer_mode") or default)
        except Exception:
            return normalize_answer_mode(default)

    def add_user_profile_note(self, user_id: str, note: str, key: str = "note") -> bool:
        uid = str(user_id or "public")[:120]
        clean = sanitize_non_instruction_context(note, limit=1500)
        if not clean:
            return False
        now = _utc_ts()
        with self._connect() as conn:
            row = conn.execute("SELECT preferences_json, notes FROM user_profiles WHERE user_id = ?", (uid,)).fetchone()
            prefs = _safe_json_loads(_row_value(row, "preferences_json", 0, "{}")) if row else {}
            notes = str(_row_value(row, "notes", 1, "") or "") if row else ""
            history = prefs.get("notes_history") or []
            if not isinstance(history, list):
                history = []
            history.append({"ts": now, "key": str(key or "note")[:80], "text": clean[:500]})
            prefs["notes_history"] = history[-30:]
            combined_notes = (notes + "\n" + clean).strip()[-4000:]
            conn.execute(
                """
                INSERT INTO user_profiles(user_id,answer_mode,preferences_json,notes,updated_at)
                VALUES (?, 'auto', ?, ?, ?)
                ON CONFLICT(user_id) DO UPDATE SET preferences_json=excluded.preferences_json, notes=excluded.notes, updated_at=excluded.updated_at
                """,
                (uid, _safe_json(prefs), combined_notes, now),
            )
        return True

    def record_answer_quality(
        self,
        *,
        interaction_id: int = 0,
        user_id: str = "public",
        channel: str = "web",
        intent: str = "general",
        answer_mode: str = "auto",
        score: float = 0,
        level: str = "",
        needs_verification: bool = False,
        verifier_model: str = "",
        verified: bool = False,
        reasons: Optional[List[str]] = None,
        metrics: Optional[Dict[str, Any]] = None,
        question: str = "",
        answer: str = "",
    ) -> int:
        with self._connect() as conn:
            cur = conn.execute(
                """
                INSERT INTO answer_quality_reports(
                    ts,interaction_id,user_id,channel,intent,answer_mode,score,level,needs_verification,
                    verifier_model,verified,reasons_json,metrics_json,question_preview,answer_preview
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    _utc_ts(), int(interaction_id or 0), str(user_id or "public")[:120], str(channel or "web")[:40],
                    str(intent or "general")[:80], normalize_answer_mode(answer_mode), float(score or 0), str(level or "")[:40],
                    1 if needs_verification else 0, str(verifier_model or "")[:160], 1 if verified else 0,
                    _safe_json(reasons or [])[:3000], _safe_json(metrics or {})[:5000],
                    sanitize_non_instruction_context(question, limit=700), sanitize_non_instruction_context(answer, limit=900),
                ),
            )
            return int(cur.lastrowid)

    def quality_dashboard(self, days: int = 14) -> Dict[str, Any]:
        since = _utc_ts() - max(1, int(days or 14)) * 86400
        with self._connect() as conn:
            total = conn.execute(
                """
                SELECT COUNT(*) AS n, COALESCE(AVG(score),0) AS avg_score,
                       COALESCE(SUM(CASE WHEN score < 0.58 THEN 1 ELSE 0 END),0) AS low_count,
                       COALESCE(SUM(CASE WHEN verified = 1 THEN 1 ELSE 0 END),0) AS verified_count
                FROM answer_quality_reports WHERE ts >= ?
                """,
                (since,),
            ).fetchone()
            by_mode = conn.execute(
                "SELECT answer_mode, COUNT(*) AS jumlah, COALESCE(AVG(score),0) AS avg_score FROM answer_quality_reports WHERE ts >= ? GROUP BY answer_mode ORDER BY jumlah DESC",
                (since,),
            ).fetchall()
            by_intent = conn.execute(
                "SELECT intent, COUNT(*) AS jumlah, COALESCE(AVG(score),0) AS avg_score FROM answer_quality_reports WHERE ts >= ? GROUP BY intent ORDER BY jumlah DESC LIMIT 15",
                (since,),
            ).fetchall()
            low = conn.execute(
                """
                SELECT id, ts, user_id, channel, intent, answer_mode, score, level, reasons_json, question_preview
                FROM answer_quality_reports WHERE ts >= ? AND score < 0.70 ORDER BY score ASC, ts DESC LIMIT 30
                """,
                (since,),
            ).fetchall()
        low_items = []
        for row in low:
            item = _row_to_dict(row)
            item["ts_wib"] = _timestamp_to_wib(float(item.get("ts") or 0))
            item["reasons"] = _safe_json_loads("{}")
            try:
                item["reasons"] = json.loads(str(item.get("reasons_json") or "[]"))
            except Exception:
                item["reasons"] = []
            low_items.append(item)
        return {
            "days": days,
            "total": int(_row_value(total, "n", 0, 0) or 0),
            "avg_score": round(float(_row_value(total, "avg_score", 1, 0) or 0), 3),
            "low_count": int(_row_value(total, "low_count", 2, 0) or 0),
            "verified_count": int(_row_value(total, "verified_count", 3, 0) or 0),
            "by_mode": [_row_to_dict(row) for row in by_mode],
            "by_intent": [_row_to_dict(row) for row in by_intent],
            "low_quality": low_items,
        }

    def weekly_quality_evaluation(self, days: int = 7, save: bool = True) -> str:
        days = max(1, int(days or 7))
        dash = self.quality_dashboard(days=days)
        learning = self.learning_dashboard(days=days)
        usage = self.usage_summary(days=days)
        gaps = self.list_knowledge_gaps(status="open", limit=10)
        lines = [
            f"📊 Evaluasi Kualitas Adioranye {days} Hari Terakhir",
            "",
            f"Total jawaban dinilai: {dash.get('total', 0)}",
            f"Rata-rata skor kualitas: {dash.get('avg_score', 0)}",
            f"Jawaban skor rendah: {dash.get('low_count', 0)}",
            f"Jawaban diverifikasi: {dash.get('verified_count', 0)}",
            f"Estimasi biaya: Rp{float(usage.get('cost_idr') or 0):,.2f}",
            "",
            "Mode jawaban:",
        ]
        for row in dash.get("by_mode", [])[:8]:
            lines.append(f"- {row.get('answer_mode')}: {row.get('jumlah')} jawaban | avg {float(row.get('avg_score') or 0):.2f}")
        lines.append("")
        lines.append("Intent paling sering:")
        for row in dash.get("by_intent", [])[:8]:
            lines.append(f"- {row.get('intent')}: {row.get('jumlah')} | avg {float(row.get('avg_score') or 0):.2f}")
        lines.append("")
        lines.append("Knowledge gap terbuka:")
        if gaps:
            for gap in gaps[:8]:
                lines.append(f"- #{gap.get('id')} [{gap.get('priority')}] {gap.get('question')[:120]}")
        else:
            lines.append("- Tidak ada gap terbuka yang tercatat.")
        lines.append("")
        lines.append("Rekomendasi:")
        if dash.get("avg_score", 0) < 0.70:
            lines.append("- Aktifkan mode riset/kritis untuk topik faktual dan tambah sumber KB prioritas.")
        if dash.get("low_count", 0):
            lines.append("- Periksa daftar jawaban skor rendah di dashboard Quality Control.")
        if gaps:
            lines.append("- Tambahkan sumber untuk knowledge gap terbuka, lalu tandai selesai.")
        if not gaps and dash.get("avg_score", 0) >= 0.75:
            lines.append("- Kualitas relatif stabil. Lanjutkan update KB sekuensial dan pantau feedback negatif.")
        report = "\n".join(lines)
        if save:
            with self._connect() as conn:
                conn.execute(
                    "INSERT INTO weekly_evaluations(ts,days,report,metrics_json) VALUES (?, ?, ?, ?)",
                    (_utc_ts(), days, report, _safe_json({"quality": dash, "learning": learning, "usage": usage})[:8000]),
                )
        return report

    def export_knowledge_base_jsonl(self, limit: int = 5000) -> str:
        with self._connect() as conn:
            rows = conn.execute(
                """
                SELECT d.id AS doc_id, d.title, d.source, d.collection, d.tags, d.created_at, d.updated_at,
                       c.chunk_index, c.heading, c.page_label, c.content
                FROM documents d LEFT JOIN chunks c ON c.doc_id = d.id
                ORDER BY d.updated_at DESC, c.chunk_index ASC LIMIT ?
                """,
                (max(1, int(limit or 5000)),),
            ).fetchall()
        lines = []
        for row in rows:
            item = _row_to_dict(row)
            lines.append(json.dumps(item, ensure_ascii=False, default=str))
        return "\n".join(lines)

    def export_interactions_jsonl(self, days: int = 30, limit: int = 5000) -> str:
        since = _utc_ts() - max(1, int(days or 30)) * 86400
        with self._connect() as conn:
            rows = conn.execute(
                "SELECT * FROM interactions WHERE ts >= ? ORDER BY ts DESC LIMIT ?",
                (since, max(1, int(limit or 5000))),
            ).fetchall()
        lines = []
        for row in rows:
            item = _row_to_dict(row)
            item["ts_wib"] = _timestamp_to_wib(float(item.get("ts") or 0))
            lines.append(json.dumps(item, ensure_ascii=False, default=str))
        return "\n".join(lines)

    def import_knowledge_base_jsonl(self, jsonl_text: str, collection_prefix: str = "Imported") -> Dict[str, Any]:
        added = 0
        skipped = 0
        errors = 0
        for line in str(jsonl_text or "").splitlines():
            line = line.strip()
            if not line:
                continue
            try:
                item = json.loads(line)
                title = str(item.get("title") or "Imported document")[:220]
                source = str(item.get("source") or "jsonl-import")[:500]
                content = str(item.get("content") or item.get("answer") or item.get("text") or "").strip()
                if not content:
                    skipped += 1
                    continue
                collection = str(item.get("collection") or collection_prefix or "Imported")[:80]
                tags = str(item.get("tags") or "import,jsonl")[:300]
                self.add_document(title=title, text=content, source=source, collection=collection, tags=tags)
                added += 1
            except Exception:
                errors += 1
        return {"added": added, "skipped": skipped, "errors": errors}

    def list_documents(self, limit: int = 20, collection: str = "") -> List[Dict[str, Any]]:
        collection_filter = _normalize_collection(collection) if str(collection or "").strip() else ""
        with self._connect() as conn:
            if collection_filter:
                rows = conn.execute(
                    """
                    SELECT d.id, d.title, d.source, d.collection, d.tags, d.pinned, d.created_at, d.updated_at,
                           COUNT(c.id) AS chunks, COALESCE(SUM(LENGTH(c.content)), 0) AS characters
                    FROM documents d LEFT JOIN chunks c ON c.doc_id = d.id
                    WHERE d.collection = ?
                    GROUP BY d.id ORDER BY d.pinned DESC, d.created_at DESC LIMIT ?
                    """,
                    (collection_filter, max(1, int(limit or 20))),
                ).fetchall()
            else:
                rows = conn.execute(
                    """
                    SELECT d.id, d.title, d.source, d.collection, d.tags, d.pinned, d.created_at, d.updated_at,
                           COUNT(c.id) AS chunks, COALESCE(SUM(LENGTH(c.content)), 0) AS characters
                    FROM documents d LEFT JOIN chunks c ON c.doc_id = d.id
                    GROUP BY d.id ORDER BY d.pinned DESC, d.created_at DESC LIMIT ?
                    """,
                    (max(1, int(limit or 20)),),
                ).fetchall()
        out = []
        for row in rows:
            item = dict(row)
            try:
                item["created_at_wib"] = datetime.fromtimestamp(float(item.get("created_at") or 0), WIB_TZ).strftime("%Y-%m-%d %H:%M:%S WIB")
            except Exception:
                item["created_at_wib"] = ""
            out.append(item)
        return out

    def knowledge_collections(self) -> List[Dict[str, Any]]:
        with self._connect() as conn:
            rows = conn.execute(
                """
                SELECT COALESCE(NULLIF(collection, ''), 'Default') AS collection,
                       COUNT(*) AS documents,
                       COALESCE(SUM((SELECT COUNT(*) FROM chunks c WHERE c.doc_id = documents.id)), 0) AS chunks
                FROM documents
                GROUP BY COALESCE(NULLIF(collection, ''), 'Default')
                ORDER BY documents DESC, collection ASC
                """
            ).fetchall()
        return [dict(row) for row in rows]

    def get_document(self, doc_id: int, max_chars: int = 8000) -> Dict[str, Any]:
        try:
            doc_id = int(doc_id)
        except Exception:
            return {}
        with self._connect() as conn:
            doc = conn.execute("SELECT * FROM documents WHERE id = ?", (doc_id,)).fetchone()
            if not doc:
                return {}
            chunks = conn.execute(
                "SELECT chunk_index, content, heading, page_label FROM chunks WHERE doc_id = ? ORDER BY chunk_index ASC",
                (doc_id,),
            ).fetchall()
        item = dict(doc)
        item["chunks"] = len(chunks)
        body_parts = []
        total = 0
        for ch in chunks:
            text = str(ch["content"] or "")
            heading = str(ch["heading"] or "").strip()
            page = str(ch["page_label"] or "").strip()
            label = f"[Chunk {ch['chunk_index']}" + (f" · {page}" if page else "") + (f" · {heading}" if heading else "") + "]"
            if total + len(text) > max_chars:
                body_parts.append("[Dokumen dipotong untuk preview]")
                break
            body_parts.append(f"{label}\n{text}")
            total += len(text)
        item["preview"] = "\n\n".join(body_parts)
        return item

    def set_document_pinned(self, doc_id: int, pinned: bool = True) -> bool:
        try:
            doc_id = int(doc_id)
        except Exception:
            return False
        with self._connect() as conn:
            cur = conn.execute("UPDATE documents SET pinned = ?, updated_at = ? WHERE id = ?", (1 if pinned else 0, _utc_ts(), doc_id))
            return bool(cur.rowcount)

    def update_document_metadata(self, doc_id: int, collection: str = "", tags: str = "") -> bool:
        try:
            doc_id = int(doc_id)
        except Exception:
            return False
        collection = _normalize_collection(collection) if collection else "Default"
        tags = _normalize_tags(tags)
        with self._connect() as conn:
            cur = conn.execute("UPDATE documents SET collection = ?, tags = ?, updated_at = ? WHERE id = ?", (collection, tags, _utc_ts(), doc_id))
            if cur.rowcount:
                self.rebuild_knowledge_index()
            return bool(cur.rowcount)

    def delete_document(self, doc_id: int) -> bool:
        try:
            doc_id = int(doc_id)
        except Exception:
            return False
        with self._connect() as conn:
            chunk_ids = [int(row["id"]) for row in conn.execute("SELECT id FROM chunks WHERE doc_id = ?", (doc_id,)).fetchall()]
            if chunk_ids:
                try:
                    conn.executemany("DELETE FROM chunks_fts WHERE rowid = ?", [(cid,) for cid in chunk_ids])
                except Exception:
                    pass
            cur = conn.execute("DELETE FROM documents WHERE id = ?", (doc_id,))
            conn.execute("DELETE FROM chunks WHERE doc_id = ?", (doc_id,))
            return bool(cur.rowcount)

    def rebuild_knowledge_index(self) -> Tuple[int, int]:
        """Rebuild FTS index from chunks. Returns (documents, chunks)."""
        with self._connect() as conn:
            try:
                conn.execute("DELETE FROM chunks_fts")
            except Exception:
                try:
                    conn.execute("DROP TABLE IF EXISTS chunks_fts")
                    conn.execute("CREATE VIRTUAL TABLE IF NOT EXISTS chunks_fts USING fts5(content, title, source, collection, tags, heading, page_label, content='')")
                except Exception:
                    pass
            rows = conn.execute(
                """
                SELECT c.id, c.content, c.heading, c.page_label, d.title, d.source, d.collection, d.tags
                FROM chunks c JOIN documents d ON d.id = c.doc_id
                ORDER BY c.id ASC
                """
            ).fetchall()
            inserted = 0
            for row in rows:
                try:
                    conn.execute(
                        "INSERT INTO chunks_fts(rowid, content, title, source, collection, tags, heading, page_label) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                        (int(row["id"]), str(row["content"] or ""), str(row["title"] or ""), str(row["source"] or ""), str(row["collection"] or "Default"), str(row["tags"] or ""), str(row["heading"] or ""), str(row["page_label"] or "")),
                    )
                    inserted += 1
                except Exception:
                    pass
            doc_count = int(conn.execute("SELECT COUNT(*) FROM documents").fetchone()[0] or 0)
            return doc_count, inserted

    def knowledge_stats(self) -> Dict[str, Any]:
        with self._connect() as conn:
            docs = int(conn.execute("SELECT COUNT(*) FROM documents").fetchone()[0] or 0)
            chunks = int(conn.execute("SELECT COUNT(*) FROM chunks").fetchone()[0] or 0)
            chars = int(conn.execute("SELECT COALESCE(SUM(LENGTH(content)), 0) FROM chunks").fetchone()[0] or 0)
            pinned = int(conn.execute("SELECT COUNT(*) FROM documents WHERE pinned = 1").fetchone()[0] or 0)
            collections = int(conn.execute("SELECT COUNT(DISTINCT COALESCE(NULLIF(collection, ''), 'Default')) FROM documents").fetchone()[0] or 0)
            latest = conn.execute("SELECT title, source, collection, created_at FROM documents ORDER BY created_at DESC LIMIT 1").fetchone()
        data = {"documents": docs, "chunks": chunks, "characters": chars, "pinned": pinned, "collections": collections}
        if latest:
            data["latest_title"] = latest["title"]
            data["latest_source"] = latest["source"]
            data["latest_collection"] = latest["collection"]
            try:
                data["latest_at_wib"] = datetime.fromtimestamp(float(latest["created_at"] or 0), WIB_TZ).strftime("%Y-%m-%d %H:%M:%S WIB")
            except Exception:
                data["latest_at_wib"] = ""
        return data


    # Usage / observability
    def log_interaction(
        self,
        user_id: str,
        channel: str,
        intent: str,
        model: str,
        question: str,
        answer: str,
        meta: Optional[Dict[str, Any]] = None,
        success: bool = True,
    ) -> int:
        meta = meta or {}
        usage = meta.get("usage") or {}
        input_tokens = int(usage.get("prompt_tokens") or usage.get("input_tokens") or 0)
        output_tokens = int(usage.get("completion_tokens") or usage.get("output_tokens") or 0)
        cost_idr = estimate_cost_idr(meta, model)
        latency = float(meta.get("latency_seconds") or meta.get("power_latency_seconds") or 0)
        quality = extract_quality_score(meta, answer)
        with self._connect() as conn:
            cur = conn.execute(
                """
                INSERT INTO interactions(ts,user_id,channel,intent,model,input_tokens,output_tokens,cost_idr,
                                         latency_seconds,success,question_preview,answer_preview,meta_json)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    _utc_ts(),
                    str(user_id or "public")[:120],
                    str(channel or "web")[:40],
                    str(intent or "")[:80],
                    str(model or "")[:160],
                    input_tokens,
                    output_tokens,
                    cost_idr,
                    latency,
                    1 if success else 0,
                    str(question or "")[:500],
                    str(answer or "")[:500],
                    _safe_json(meta)[:12000],
                ),
            )
            interaction_id = int(cur.lastrowid or 0)
        self.update_model_score(model=model, intent=intent, success=success, latency_seconds=latency, quality_score=quality, cost_idr=cost_idr)
        if success:
            self.register_model_success(model)
        else:
            self.register_model_failure(model, error=str((meta or {}).get("error") or "failed_interaction"))
        return interaction_id

    def usage_summary(self, days: int = 1) -> Dict[str, Any]:
        since = _utc_ts() - max(1, int(days or 1)) * 86400
        by_model_columns = ["model", "requests", "input_tokens", "output_tokens", "cost_idr", "avg_latency"]
        with self._connect() as conn:
            rows = conn.execute(
                """
                SELECT model, COUNT(*) AS requests, COALESCE(SUM(input_tokens), 0) AS input_tokens,
                       COALESCE(SUM(output_tokens), 0) AS output_tokens, COALESCE(SUM(cost_idr), 0) AS cost_idr,
                       COALESCE(AVG(latency_seconds), 0) AS avg_latency
                FROM interactions WHERE ts >= ? GROUP BY model ORDER BY cost_idr DESC
                """,
                (since,),
            ).fetchall()
            total = conn.execute(
                "SELECT COUNT(*) AS requests, COALESCE(SUM(cost_idr), 0) AS cost_idr FROM interactions WHERE ts >= ?",
                (since,),
            ).fetchone()
        return {
            "days": int(days or 1),
            "requests": int(_row_value(total, "requests", 0, 0) or 0),
            "cost_idr": float(_row_value(total, "cost_idr", 1, 0) or 0),
            "by_model": [_row_to_dict(row, by_model_columns) for row in rows],
        }

    def count_expensive_calls_today(self) -> int:
        start = datetime.now(WIB_TZ).replace(hour=0, minute=0, second=0, microsecond=0).timestamp()
        with self._connect() as conn:
            rows = conn.execute("SELECT model FROM interactions WHERE ts >= ?", (start,)).fetchall()
        return sum(
            1
            for row in rows
            if model_cost_tier(str(_row_value(row, "model", 0, "") or "")) in {"medium", "expensive", "ultra"}
        )

    def database_overview(self) -> Dict[str, Any]:
        """Return compact DB statistics for admin health center."""
        with self._connect() as conn:
            tables = {
                "memories": "memories",
                "documents": "documents",
                "chunks": "chunks",
                "interactions": "interactions",
                "benchmarks": "benchmarks",
                "response_cache": "response_cache",
                "model_scores": "model_scores",
                "circuit_breakers": "circuit_breakers",
            }
            out: Dict[str, Any] = {}
            for key, table in tables.items():
                try:
                    out[key] = int(conn.execute(f"SELECT COUNT(*) FROM {table}").fetchone()[0] or 0)
                except Exception:
                    out[key] = 0
        try:
            size = Path(self.db_path).stat().st_size
            units = ["B", "KB", "MB", "GB"]
            value = float(size)
            for unit in units:
                if value < 1024 or unit == units[-1]:
                    out["db_size"] = f"{value:.1f} {unit}"
                    break
                value /= 1024
        except Exception:
            out["db_size"] = "0 B"
        return out

    def cleanup_old_data(self, log_days: int = 30, cache_days: int = 7, benchmark_days: int = 14) -> Dict[str, int]:
        """Delete old operational rows while preserving knowledge base and memories."""
        now = _utc_ts()
        deleted: Dict[str, int] = {}
        with self._connect() as conn:
            cur = conn.execute("DELETE FROM interactions WHERE ts < ?", (now - max(1, int(log_days or 30)) * 86400,))
            deleted["interactions"] = int(cur.rowcount or 0)
            cur = conn.execute("DELETE FROM response_cache WHERE expires_at < ? OR ts < ?", (now, now - max(1, int(cache_days or 7)) * 86400))
            deleted["response_cache"] = int(cur.rowcount or 0)
            cur = conn.execute("DELETE FROM benchmarks WHERE ts < ?", (now - max(1, int(benchmark_days or 14)) * 86400,))
            deleted["benchmarks"] = int(cur.rowcount or 0)
            try:
                conn.execute("VACUUM")
            except Exception:
                pass
        return deleted

    def clear_usage_logs(self) -> int:
        with self._connect() as conn:
            cur = conn.execute("DELETE FROM interactions")
            return int(cur.rowcount or 0)

    def clear_response_cache(self) -> int:
        with self._connect() as conn:
            cur = conn.execute("DELETE FROM response_cache")
            return int(cur.rowcount or 0)

    def clear_memories_all(self) -> int:
        with self._connect() as conn:
            cur = conn.execute("DELETE FROM memories")
            return int(cur.rowcount or 0)

    def clear_knowledge_base(self) -> Dict[str, int]:
        with self._connect() as conn:
            counts = {
                "documents": int(conn.execute("SELECT COUNT(*) FROM documents").fetchone()[0] or 0),
                "chunks": int(conn.execute("SELECT COUNT(*) FROM chunks").fetchone()[0] or 0),
            }
            conn.execute("DELETE FROM chunks")
            conn.execute("DELETE FROM documents")
            try:
                conn.execute("DELETE FROM chunks_fts")
            except Exception:
                pass
            return counts


    # Persistent response cache
    def get_cached_response(self, cache_key: str) -> Optional[Tuple[str, Dict[str, Any]]]:
        """Return cached answer if present and not expired."""
        key = str(cache_key or "").strip()
        if not key:
            return None
        now = _utc_ts()
        with self._connect() as conn:
            row = conn.execute(
                "SELECT answer, meta_json, expires_at FROM response_cache WHERE cache_key = ? LIMIT 1",
                (key,),
            ).fetchone()
            if not row:
                return None
            expires_at = float(_row_value(row, "expires_at", 2, 0) or 0)
            if expires_at and expires_at < now:
                conn.execute("DELETE FROM response_cache WHERE cache_key = ?", (key,))
                return None
            answer = str(_row_value(row, "answer", 0, "") or "")
            meta = _safe_json_loads(_row_value(row, "meta_json", 1, "{}"))
            meta["power_response_cache_hit"] = True
            return answer, meta

    def set_cached_response(self, cache_key: str, answer: str, meta: Optional[Dict[str, Any]] = None, ttl_seconds: int = 1800) -> None:
        """Store a response cache row in SQLite."""
        key = str(cache_key or "").strip()
        body = str(answer or "").strip()
        if not key or not body:
            return
        meta = meta or {}
        now = _utc_ts()
        expires_at = now + max(60, int(ttl_seconds or 1800))
        model = str(meta.get("active_model_final") or meta.get("model_requested") or meta.get("model") or "")[:180]
        intent = str(meta.get("power_intent") or meta.get("intent") or "")[:80]
        with self._connect() as conn:
            conn.execute(
                """
                INSERT OR REPLACE INTO response_cache(cache_key,ts,expires_at,model,intent,answer,meta_json)
                VALUES (?, ?, ?, ?, ?, ?, ?)
                """,
                (key, now, expires_at, model, intent, body, _safe_json(meta)[:12000]),
            )

    # Performance optimizer: semantic cache, retrieval eval, observability, maintenance
    def _semantic_terms(self, text: str) -> str:
        try:
            return " ".join(_tokenize(text)[:80])
        except Exception:
            return str(text or "")[:500]

    def get_semantic_cached_response(
        self,
        question: str,
        intent: str = "",
        threshold: float = 0.76,
        ttl_seconds: int = 86400,
        user_id: str = "public",
        channel: str = "web",
    ) -> Optional[Tuple[str, Dict[str, Any]]]:
        """Return a cached answer for semantically similar questions.

        This is intentionally lightweight: it compares recent cached questions by
        token overlap, so it works on Streamlit Cloud without embedding services.
        """
        q = str(question or "").strip()
        if not q or len(q) < 12:
            return None
        now = _utc_ts()
        threshold = max(0.35, min(0.95, float(threshold or 0.76)))
        with self._connect() as conn:
            rows = conn.execute(
                """
                SELECT id, question, answer, meta_json, expires_at, intent, user_id, channel
                FROM semantic_response_cache
                WHERE expires_at > ? AND (? = '' OR intent = ? OR intent = '')
                ORDER BY ts DESC LIMIT 180
                """,
                (now, str(intent or "")[:80], str(intent or "")[:80]),
            ).fetchall()
        best: Optional[Tuple[float, Any]] = None
        for row in rows:
            candidate_q = str(_row_value(row, "question", 1, "") or "")
            sim = semantic_similarity(q, candidate_q)
            if sim >= threshold and (best is None or sim > best[0]):
                best = (sim, row)
        if not best:
            return None
        sim, row = best
        answer = str(_row_value(row, "answer", 2, "") or "")
        meta = _safe_json_loads(_row_value(row, "meta_json", 3, "{}"))
        meta["semantic_cache_hit"] = True
        meta["semantic_cache_similarity"] = round(float(sim), 4)
        meta["semantic_cache_question"] = str(_row_value(row, "question", 1, "") or "")[:500]
        meta["power_response_cache_hit"] = True
        return answer, meta

    def set_semantic_cached_response(
        self,
        question: str,
        answer: str,
        meta: Optional[Dict[str, Any]] = None,
        intent: str = "",
        ttl_seconds: int = 86400,
        user_id: str = "public",
        channel: str = "web",
    ) -> None:
        q = str(question or "").strip()
        body = str(answer or "").strip()
        if not q or not body or len(q) < 10:
            return
        meta = meta or {}
        now = _utc_ts()
        expires_at = now + max(600, int(ttl_seconds or 86400))
        with self._connect() as conn:
            conn.execute(
                """
                INSERT INTO semantic_response_cache(ts, expires_at, user_id, channel, intent, question, question_terms, answer, meta_json)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (now, expires_at, str(user_id or "public")[:120], str(channel or "web")[:40], str(intent or "")[:80], q[:2000], self._semantic_terms(q), body, _safe_json(meta)[:12000]),
            )
            # Keep the table compact.
            conn.execute(
                "DELETE FROM semantic_response_cache WHERE expires_at < ? OR id NOT IN (SELECT id FROM semantic_response_cache ORDER BY ts DESC LIMIT 1000)",
                (now,),
            )

    def record_retrieval_eval(
        self,
        question: str,
        search_query: str,
        sources: List[Dict[str, Any]],
        intent: str = "",
        user_id: str = "public",
        channel: str = "web",
        latency_seconds: float = 0.0,
        meta: Optional[Dict[str, Any]] = None,
    ) -> int:
        metrics = retrieval_precision_estimate(search_query or question, sources or [])
        now = _utc_ts()
        with self._connect() as conn:
            cur = conn.execute(
                """
                INSERT INTO retrieval_eval_reports(ts, user_id, channel, intent, question, search_query, source_count, precision_estimate, avg_similarity, latency_seconds, meta_json)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    now,
                    str(user_id or "public")[:120],
                    str(channel or "web")[:40],
                    str(intent or "")[:80],
                    str(question or "")[:2000],
                    str(search_query or "")[:2000],
                    int(metrics.get("source_count") or 0),
                    float(metrics.get("precision") or 0),
                    float(metrics.get("avg_similarity") or 0),
                    float(latency_seconds or 0),
                    _safe_json({"metrics": metrics, **(meta or {})})[:12000],
                ),
            )
            return int(cur.lastrowid or 0)

    def record_source_runtime(self, source: str, success: bool = True, latency_seconds: float = 0.0, quality: float = 0.0, error: str = "") -> None:
        src = str(source or "").strip()[:600]
        if not src:
            return
        now = _utc_ts()
        with self._connect() as conn:
            conn.execute(
                """
                INSERT INTO source_runtime_stats(source,total_requests,success_count,failure_count,total_latency,total_quality,last_error,last_seen_at,status,updated_at)
                VALUES (?, 1, ?, ?, ?, ?, ?, ?, 'active', ?)
                ON CONFLICT(source) DO UPDATE SET
                    total_requests = total_requests + 1,
                    success_count = success_count + excluded.success_count,
                    failure_count = failure_count + excluded.failure_count,
                    total_latency = total_latency + excluded.total_latency,
                    total_quality = total_quality + excluded.total_quality,
                    last_error = excluded.last_error,
                    last_seen_at = excluded.last_seen_at,
                    updated_at = excluded.updated_at
                """,
                (src, 1 if success else 0, 0 if success else 1, float(latency_seconds or 0), float(quality or 0), str(error or "")[:1000], now, now),
            )

    def performance_dashboard(self, days: int = 14) -> Dict[str, Any]:
        since = _utc_ts() - max(1, int(days or 14)) * 86400
        out: Dict[str, Any] = {"days": days}
        with self._connect() as conn:
            row = conn.execute(
                """
                SELECT COUNT(*) AS total, COALESCE(AVG(precision_estimate),0) AS avg_precision,
                       COALESCE(AVG(avg_similarity),0) AS avg_similarity,
                       COALESCE(AVG(latency_seconds),0) AS avg_retrieval_latency
                FROM retrieval_eval_reports WHERE ts >= ?
                """,
                (since,),
            ).fetchone()
            out["retrieval"] = _row_to_dict(row)
            row = conn.execute(
                "SELECT COUNT(*) AS total FROM semantic_response_cache WHERE expires_at > ?",
                (_utc_ts(),),
            ).fetchone()
            out["semantic_cache_active"] = int(_row_value(row, "total", 0, 0) or 0)
            out["semantic_cache_hits"] = int(conn.execute(
                "SELECT COUNT(*) FROM interactions WHERE ts >= ? AND meta_json LIKE '%semantic_cache_hit%'",
                (since,),
            ).fetchone()[0] or 0)
            out["top_intents_latency"] = [dict(r) for r in conn.execute(
                """
                SELECT intent, COUNT(*) AS jumlah, ROUND(COALESCE(AVG(latency_seconds),0),3) AS avg_latency
                FROM interactions WHERE ts >= ? GROUP BY intent ORDER BY jumlah DESC LIMIT 12
                """,
                (since,),
            ).fetchall()]
            out["slow_sources"] = [dict(r) for r in conn.execute(
                """
                SELECT source, total_requests, failure_count,
                       ROUND(total_latency / CASE WHEN total_requests=0 THEN 1 ELSE total_requests END, 3) AS avg_latency,
                       status
                FROM source_runtime_stats
                ORDER BY avg_latency DESC, failure_count DESC LIMIT 20
                """
            ).fetchall()]
            out["recent_retrieval"] = [dict(r) for r in conn.execute(
                """
                SELECT datetime(ts, 'unixepoch', 'localtime') AS waktu, intent, source_count, precision_estimate, avg_similarity,
                       substr(question,1,120) AS question
                FROM retrieval_eval_reports ORDER BY ts DESC LIMIT 30
                """
            ).fetchall()]
        try:
            overview = self.database_overview()
            out["database"] = overview
        except Exception:
            out["database"] = {}
        return out

    def optimize_database(self, vacuum: bool = False) -> Dict[str, Any]:
        result: Dict[str, Any] = {"ok": True, "vacuum": bool(vacuum)}
        started = time.time()
        try:
            with self._connect() as conn:
                try:
                    conn.execute("PRAGMA optimize")
                except Exception:
                    pass
                try:
                    conn.execute("ANALYZE")
                except Exception:
                    pass
                if vacuum:
                    try:
                        conn.execute("VACUUM")
                    except Exception as exc:
                        result["vacuum_error"] = str(exc)[:500]
                try:
                    result["integrity_check"] = str(conn.execute("PRAGMA quick_check").fetchone()[0])
                except Exception as exc:
                    result["integrity_error"] = str(exc)[:500]
        except Exception as exc:
            result["ok"] = False
            result["error"] = str(exc)[:1000]
        result["latency_seconds"] = round(time.time() - started, 3)
        return result

    # Adaptive scoring and circuit breaker
    def update_model_score(
        self,
        model: str,
        intent: str = "general",
        success: bool = True,
        latency_seconds: float = 0,
        quality_score: float = 0,
        cost_idr: float = 0,
    ) -> None:
        """Accumulate model performance per intent for adaptive routing."""
        model_name = str(model or "").strip()
        if not model_name:
            return
        intent_name = str(intent or "general").strip()[:80] or "general"
        now = _utc_ts()
        ok = 1 if success else 0
        err = 0 if success else 1
        with self._connect() as conn:
            conn.execute(
                """
                INSERT INTO model_scores(
                    model,intent,total_requests,success_count,error_count,total_latency,total_quality,total_cost,
                    last_success_at,last_error_at,updated_at
                ) VALUES (?, ?, 1, ?, ?, ?, ?, ?, ?, ?, ?)
                ON CONFLICT(model,intent) DO UPDATE SET
                    total_requests = total_requests + 1,
                    success_count = success_count + excluded.success_count,
                    error_count = error_count + excluded.error_count,
                    total_latency = total_latency + excluded.total_latency,
                    total_quality = total_quality + excluded.total_quality,
                    total_cost = total_cost + excluded.total_cost,
                    last_success_at = CASE WHEN excluded.success_count > 0 THEN excluded.last_success_at ELSE last_success_at END,
                    last_error_at = CASE WHEN excluded.error_count > 0 THEN excluded.last_error_at ELSE last_error_at END,
                    updated_at = excluded.updated_at
                """,
                (
                    model_name,
                    intent_name,
                    ok,
                    err,
                    max(0.0, float(latency_seconds or 0)),
                    max(0.0, min(1.0, float(quality_score or 0))),
                    max(0.0, float(cost_idr or 0)),
                    now if success else 0,
                    now if not success else 0,
                    now,
                ),
            )

    def register_model_success(self, model: str) -> None:
        """Close/reset circuit breaker after a successful response."""
        model_name = str(model or "").strip()
        if not model_name:
            return
        with self._connect() as conn:
            conn.execute(
                """
                INSERT INTO circuit_breakers(model,failure_count,open_until,last_error,updated_at)
                VALUES (?, 0, 0, '', ?)
                ON CONFLICT(model) DO UPDATE SET
                    failure_count = 0,
                    open_until = 0,
                    last_error = '',
                    updated_at = excluded.updated_at
                """,
                (model_name, _utc_ts()),
            )

    def register_model_failure(
        self,
        model: str,
        error: str = "",
        max_failures: int = 3,
        cooldown_seconds: int = 1800,
    ) -> None:
        """Increase model failure count and open the circuit if failures pass the threshold."""
        model_name = str(model or "").strip()
        if not model_name:
            return
        now = _utc_ts()
        max_failures = max(1, int(max_failures or 3))
        cooldown = max(60, int(cooldown_seconds or 1800))
        with self._connect() as conn:
            row = conn.execute(
                "SELECT failure_count FROM circuit_breakers WHERE model = ? LIMIT 1",
                (model_name,),
            ).fetchone()
            next_failures = int(_row_value(row, "failure_count", 0, 0) or 0) + 1
            open_until = now + cooldown if next_failures >= max_failures else 0
            conn.execute(
                """
                INSERT INTO circuit_breakers(model,failure_count,open_until,last_error,updated_at)
                VALUES (?, ?, ?, ?, ?)
                ON CONFLICT(model) DO UPDATE SET
                    failure_count = excluded.failure_count,
                    open_until = excluded.open_until,
                    last_error = excluded.last_error,
                    updated_at = excluded.updated_at
                """,
                (model_name, next_failures, open_until, str(error or "")[:1000], now),
            )

    def is_model_blocked(self, model: str) -> bool:
        model_name = str(model or "").strip()
        if not model_name:
            return False
        now = _utc_ts()
        with self._connect() as conn:
            row = conn.execute(
                "SELECT open_until FROM circuit_breakers WHERE model = ? LIMIT 1",
                (model_name,),
            ).fetchone()
        return bool(float(_row_value(row, "open_until", 0, 0) or 0) > now)

    def filter_blocked_models(self, models: List[str]) -> List[str]:
        """Remove models with open circuit breaker. If all are blocked, return original list so errors remain visible."""
        original = [str(m).strip() for m in (models or []) if str(m or "").strip()]
        if not original:
            return []
        now = _utc_ts()
        with self._connect() as conn:
            rows = conn.execute("SELECT model, open_until FROM circuit_breakers WHERE open_until > ?", (now,)).fetchall()
        blocked = {str(_row_value(row, "model", 0, "")) for row in rows}
        filtered = [m for m in original if m not in blocked]
        return filtered or original

    def _computed_model_score(self, row: Dict[str, Any], model_name: str = "") -> float:
        total = max(1, int(row.get("total_requests") or 0))
        success_rate = float(row.get("success_count") or 0) / total
        error_rate = float(row.get("error_count") or 0) / total
        avg_latency = float(row.get("total_latency") or 0) / total
        avg_quality = float(row.get("total_quality") or 0) / total
        avg_cost = float(row.get("total_cost") or 0) / total
        tier = model_cost_tier(model_name or str(row.get("model") or ""))
        tier_bonus = {"cheap": 0.08, "medium": 0.03, "expensive": 0.0, "ultra": -0.12}.get(tier, -0.03)
        latency_penalty = min(avg_latency / 30.0, 0.35)
        cost_penalty = min(avg_cost / 50.0, 0.25)
        score = (success_rate * 0.42) + (avg_quality * 0.32) + tier_bonus - (error_rate * 0.25) - latency_penalty - cost_penalty
        if total < 3:
            score -= 0.04
        return round(max(0.0, min(1.0, score)), 4)

    def rank_models_for_intent(self, models: List[str], intent: str = "general") -> List[str]:
        """Rank candidate models using per-intent history, fallbacks, cost tier, and circuit breaker state."""
        candidates = [str(m).strip() for m in (models or []) if str(m or "").strip()]
        candidates = list(dict.fromkeys(candidates))
        if not candidates:
            return []
        intent_name = str(intent or "general").strip()[:80] or "general"
        unblocked = self.filter_blocked_models(candidates)
        now = _utc_ts()
        with self._connect() as conn:
            rows = conn.execute(
                "SELECT * FROM model_scores WHERE model IN (%s) AND intent IN (?, 'general')" % ",".join("?" for _ in candidates),
                tuple(candidates) + (intent_name,),
            ).fetchall()
        by_key: Dict[Tuple[str, str], Dict[str, Any]] = {}
        for row in rows:
            item = _row_to_dict(row)
            by_key[(str(item.get("model") or ""), str(item.get("intent") or ""))] = item

        def cold_start_score(model_name: str) -> float:
            tier = model_cost_tier(model_name)
            price = model_price(model_name)
            out_price = int(price.get("output") or 999999)
            # Intent hints; these help before enough history exists.
            lower = model_name.lower()
            score = 0.50
            if tier == "cheap":
                score += 0.10
            elif tier == "medium":
                score += 0.05
            elif tier == "ultra":
                score -= 0.18
            if intent_name in {"coding"} and any(k in lower for k in ["coder", "codex", "deepseek", "qwen"]):
                score += 0.13
            if intent_name in {"academic", "research", "deep_reasoning", "document_question"} and any(k in lower for k in ["sonnet", "qwen", "gpt-5.2", "gpt-5.4", "gpt-5.5", "glm", "kimi"]):
                score += 0.10
            if intent_name in {"quick_chat", "creative", "general"} and any(k in lower for k in ["nano", "mini", "flash", "haiku", "fast"]):
                score += 0.08
            score -= min(out_price / 25000.0, 1.0) * 0.08
            return round(max(0.0, min(1.0, score)), 4)

        def score_candidate(model_name: str) -> Tuple[float, int, int, str]:
            blocked_penalty = -1.0 if model_name not in unblocked else 0.0
            specific = by_key.get((model_name, intent_name))
            general = by_key.get((model_name, "general"))
            if specific:
                score = self._computed_model_score(specific, model_name)
            elif general:
                score = self._computed_model_score(general, model_name) * 0.92
            else:
                score = cold_start_score(model_name)
            tier = model_cost_tier(model_name)
            tier_rank = {"cheap": 0, "medium": 1, "expensive": 2, "ultra": 3}.get(tier, 4)
            output_price = int(model_price(model_name).get("output") or 999999999)
            return (score + blocked_penalty, -tier_rank, -output_price, model_name)

        # Sort descending score while keeping deterministic tie-breakers.
        return sorted(candidates, key=score_candidate, reverse=True)

    def model_score_rows(self, intent: Optional[str] = None, limit: int = 100) -> List[Dict[str, Any]]:
        """Rows for admin/Telegram adaptive model score dashboard."""
        params: Tuple[Any, ...]
        where = ""
        if intent:
            where = "WHERE intent = ?"
            params = (str(intent)[:80], max(1, int(limit or 100)))
        else:
            params = (max(1, int(limit or 100)),)
        with self._connect() as conn:
            rows = conn.execute(
                f"SELECT * FROM model_scores {where} ORDER BY updated_at DESC LIMIT ?",
                params,
            ).fetchall()
        out: List[Dict[str, Any]] = []
        for row in rows:
            item = _row_to_dict(row)
            total = max(1, int(item.get("total_requests") or 0))
            item["success_rate"] = round(float(item.get("success_count") or 0) / total, 3)
            item["error_rate"] = round(float(item.get("error_count") or 0) / total, 3)
            item["avg_latency"] = round(float(item.get("total_latency") or 0) / total, 3)
            item["avg_quality"] = round(float(item.get("total_quality") or 0) / total, 3)
            item["avg_cost"] = round(float(item.get("total_cost") or 0) / total, 4)
            item["computed_score"] = self._computed_model_score(item, str(item.get("model") or ""))
            item["tier"] = model_cost_tier(str(item.get("model") or ""))
            item["last_success_wib"] = _timestamp_to_wib(float(item.get("last_success_at") or 0)) if float(item.get("last_success_at") or 0) else ""
            item["last_error_wib"] = _timestamp_to_wib(float(item.get("last_error_at") or 0)) if float(item.get("last_error_at") or 0) else ""
            item["updated_wib"] = _timestamp_to_wib(float(item.get("updated_at") or 0)) if float(item.get("updated_at") or 0) else ""
            out.append(item)
        out.sort(key=lambda x: float(x.get("computed_score") or 0), reverse=True)
        return out[: max(1, int(limit or 100))]

    def circuit_breaker_status(self, limit: int = 100) -> List[Dict[str, Any]]:
        """Return circuit breaker rows for dashboard/Telegram."""
        now = _utc_ts()
        with self._connect() as conn:
            rows = conn.execute(
                "SELECT * FROM circuit_breakers ORDER BY open_until DESC, failure_count DESC, updated_at DESC LIMIT ?",
                (max(1, int(limit or 100)),),
            ).fetchall()
        out: List[Dict[str, Any]] = []
        for row in rows:
            item = _row_to_dict(row)
            open_until = float(item.get("open_until") or 0)
            item["blocked"] = bool(open_until > now)
            item["open_until_wib"] = _timestamp_to_wib(open_until) if open_until else ""
            item["updated_wib"] = _timestamp_to_wib(float(item.get("updated_at") or 0)) if float(item.get("updated_at") or 0) else ""
            item["tier"] = model_cost_tier(str(item.get("model") or ""))
            item["last_error"] = str(item.get("last_error") or "")[:180]
            out.append(item)
        return out

    def add_benchmark(self, model: str, task: str, score: float, latency_seconds: float, success: bool, error: str = "", meta: Optional[Dict[str, Any]] = None) -> None:
        with self._connect() as conn:
            conn.execute(
                "INSERT INTO benchmarks(ts,model,task,score,latency_seconds,success,error,meta_json) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                (_utc_ts(), model, task, float(score or 0), float(latency_seconds or 0), 1 if success else 0, str(error or "")[:1000], _safe_json(meta or {})[:4000]),
            )
        self.update_model_score(model=model, intent=task, success=success, latency_seconds=latency_seconds, quality_score=score, cost_idr=0)
        if success:
            self.register_model_success(model)
        else:
            self.register_model_failure(model, error=error)

    def latest_benchmarks(self, limit: int = 50) -> List[Dict[str, Any]]:
        columns = ["id", "ts", "model", "task", "score", "latency_seconds", "success", "error", "meta_json"]
        with self._connect() as conn:
            rows = conn.execute("SELECT * FROM benchmarks ORDER BY ts DESC LIMIT ?", (max(1, int(limit or 50)),)).fetchall()
        return [_row_to_dict(row, columns) for row in rows]


_STORE_CACHE: Dict[str, PowerStore] = {}


def get_power_store(db_path: str = DEFAULT_POWER_DB) -> PowerStore:
    path = str(db_path or DEFAULT_POWER_DB)
    if path not in _STORE_CACHE:
        _STORE_CACHE[path] = PowerStore(path)
    return _STORE_CACHE[path]


# =========================
# Intent & prompt policy
# =========================

def classify_intent_text(text: str) -> str:
    t = str(text or "").lower().strip()
    wc = len(t.split())
    if not t:
        return "empty"
    if t.startswith(("/", "!")):
        return "admin_command"
    if is_music_chart_query(t):
        return "music_chart"
    if any(x in t for x in ["```", "def ", "class ", "traceback", "error", "bug", "streamlit", "api", "vercel", "github", "kode", "coding", "python", "javascript", "patch"]):
        return "coding"
    try:
        risk = detect_high_risk_question(t, intent="")
        if risk.get("is_high_risk") and any(x in t for x in ["terbaru", "terkini", "saat ini", "hari ini", "update", "apakah benar", "benarkah", "hoaks", "hoax", "valid", "bukti"]):
            return "critical_current"
    except Exception:
        pass
    if any(x in t for x in ["peternakan", "ternak", "unggas", "sapi", "kambing", "pakan", "pmk", "rabies", "flu burung", "veteriner", "hewan"]):
        return "livestock"
    if any(x in t for x in ["kesehatan", "medis", "penyakit", "obat", "gejala", "diagnosis", "terapi", "klinis", "rumah sakit"]):
        return "health"
    if any(x in t for x in ["skripsi", "jurnal", "quartile", "q1", "q2", "q3", "q4", "sinta", "scopus", "bab ", "metode", "kutipan", "referensi", "smartpls", "penelitian", "akademik", "tesis"]):
        return "academic"
    if any(x in t for x in ["hitung", "rumus", "berapa", "persentase", "kalkulasi", "calculate"]):
        return "calculation"
    if any(x in t for x in ["dokumen", "file", "pdf", "rag", "knowledge", "sumber", "berdasarkan file", "berdasarkan dokumen"]):
        return "document_question"
    if any(x in t for x in ["riset", "cari data", "terbaru", "berita", "validasi", "cek sumber", "jurnal terbaru"]):
        return "research"
    if any(x in t for x in ["caption", "konten", "desain", "copywriting", "promosi", "judul produk", "iklan"]):
        return "creative"
    if any(x in t for x in ["analisis", "analisa", "evaluasi", "strategi", "arsitektur", "bandingkan", "solusi terbaik", "algoritma", "optimasi"]):
        return "deep_reasoning"
    if wc <= 12:
        return "quick_chat"
    return "general"


PROMPT_TEMPLATES: Dict[str, str] = {
    "coding": "Jawab sebagai code reviewer senior. Berikan diagnosis, letak masalah, patch/kode yang bisa ditempel, dan langkah test singkat.",
    "academic": "Jawab dengan struktur akademik yang rapi, bahasa natural, tidak mengarang sumber, dan berikan poin yang langsung bisa dipakai.",
    "calculation": "Hitung secara teliti, tampilkan rumus/angka utama, lalu berikan jawaban akhir yang jelas.",
    "document_question": "Jawab berdasarkan konteks dokumen/knowledge base yang tersedia. Jika tidak ada sumber yang cukup, katakan keterbatasannya.",
    "research": "Pisahkan fakta, asumsi, dan langkah verifikasi. Jangan mengarang data terbaru.",
    "creative": "Berikan output kreatif yang siap pakai, ringkas, dan sesuai konteks komersial/branding.",
    "deep_reasoning": "Analisis masalah secara bertahap, prioritaskan solusi praktis, dan berikan rekomendasi akhir yang jelas.",
    "livestock": "Jawab sebagai asisten pengetahuan peternakan. Prioritaskan sumber resmi, jurnal, kesehatan hewan, pakan, manajemen ternak, dan beri catatan kehati-hatian bila data belum pasti.",
    "health": "Jawab dengan kehati-hatian kesehatan: edukatif, tidak menggantikan tenaga medis, prioritaskan sumber resmi, dan sarankan pemeriksaan profesional untuk kondisi serius.",
    "critical_current": "Jawab sebagai analis isu terkini. Utamakan data terbaru, kualitas sumber, perbedaan klaim, batas kepastian, dan rekomendasi verifikasi.",
    "music_chart": "Jawab ringkas sebagai asisten hiburan/musik. Untuk tangga lagu, gunakan konteks chart yang tersedia, sebutkan bahwa peringkat bersifat snapshot dan bisa berubah. Jangan mengarang judul/artis di luar konteks.",
}


def enhance_prompt_for_intent(user_text: str, intent: str, enable_templates: bool = True) -> str:
    if not enable_templates:
        return user_text
    template = PROMPT_TEMPLATES.get(intent)
    critical_instruction = build_critical_answer_instruction(user_text)
    if not template and not critical_instruction:
        return user_text
    parts = [str(user_text or "")]
    if template:
        parts.append(f"Instruksi mode {intent}: {template}")
    if critical_instruction:
        parts.append(critical_instruction)
    return "\n\n".join(parts)


def build_power_context(
    store: PowerStore,
    user_text: str,
    base_memory: str = "",
    user_id: str = "global",
    enable_rag: bool = True,
    rag_top_k: int = 5,
    enable_persistent_memory: bool = True,
    memory_top_k: int = 8,
    retrieval_query: str = "",
    preselected_docs: Optional[List[Dict[str, Any]]] = None,
) -> str:
    sections: List[str] = []
    base = sanitize_non_instruction_context(base_memory, limit=3500)
    if base:
        sections.append(base)
    if enable_persistent_memory:
        memories = store.search_memories(user_text, user_id=user_id, limit=memory_top_k)
        if memories:
            lines = [f"- {sanitize_non_instruction_context(m.get('text', ''), limit=500)}" for m in memories if str(m.get("text", "")).strip()]
            if lines:
                sections.append("MEMORY SQLITE RELEVAN (konteks non-instruksi):\n" + "\n".join(lines)[:3000])
    try:
        templates = store.search_answer_templates(user_text, intent=classify_intent_text(user_text), limit=2)
        if templates:
            t_lines = []
            for tmpl in templates:
                preview = sanitize_non_instruction_context(str(tmpl.get("answer") or ""), limit=700)
                if preview:
                    t_lines.append(f"- Template {tmpl.get('title')}: {preview}")
            if t_lines:
                sections.append("TEMPLATE JAWABAN RELEVAN (referensi gaya/struktur, bukan instruksi mutlak):\n" + "\n".join(t_lines))
    except Exception:
        pass
    retrieval_q = str(retrieval_query or user_text or "")
    try:
        critical_detection = detect_critical_question(user_text)
        if critical_detection.get("is_critical"):
            claims = store.search_current_claims(retrieval_q, limit=6, days=180)
            if claims:
                c_lines = []
                for idx, claim in enumerate(claims, start=1):
                    body = re.sub(r"\s+", " ", sanitize_non_instruction_context(str(claim.get("claim") or ""), limit=650)).strip()
                    c_lines.append(f"[CLAIM{idx}] {body} (sumber {claim.get('source_name') or claim.get('source')}; kualitas {claim.get('source_quality')}; freshness {claim.get('freshness_score')}; tanggal {claim.get('published_wib') or claim.get('ts_wib')})")
                sections.append("KLAIM/FAKTA TERKINI DARI CURRENT CRITICAL LAYER (konteks non-instruksi):\n" + "\n".join(c_lines))
    except Exception:
        pass

    if enable_rag:
        docs = list(preselected_docs or [])
        if not docs:
            docs = store.search_documents(retrieval_q, limit=rag_top_k)
        if docs:
            lines = []
            for idx, doc in enumerate(docs, start=1):
                content = re.sub(r"\s+", " ", sanitize_non_instruction_context(str(doc.get("content", "")), limit=1200)).strip()
                lines.append(f"[KB{idx}] {doc.get('title')} (chunk {doc.get('chunk_index')}, kualitas sumber {doc.get('source_quality', 55)}, sumber {doc.get('source')}): {content}")
            sections.append("KONTEKS KNOWLEDGE BASE/RAG NON-INSTRUKSI:\n" + "\n\n".join(lines))
    return "\n\n".join([s for s in sections if s.strip()])[:10000]



CASUAL_CHAT_PATTERNS = {
    "halo", "hai", "hi", "hello", "pagi", "siang", "sore", "malam",
    "apa kabar", "terima kasih", "makasih", "thanks", "thank you", "oke", "ok",
    "siap", "mantap", "lanjut", "iya", "tidak", "tes", "test",
}

SOURCE_REQUEST_TERMS = [
    "sumber", "source", "referensi", "rujukan", "bukti", "validasi", "cek fakta",
    "berdasarkan kb", "berdasarkan knowledge", "berdasarkan dokumen", "berdasarkan file",
    "kutipan", "sitasi", "citation", "jurnal", "q1", "q2", "q3", "q4", "sinta", "scopus",
]

SOURCE_WORTHY_INTENTS = {
    "research", "academic", "health", "livestock", "critical_current", "document_question", "music_chart",
}


def is_casual_or_light_chat(user_text: str, intent: str = "", answer_mode: str = "") -> bool:
    """Return True for greetings/light conversation that should not use or display KB sources."""
    text = re.sub(r"\s+", " ", str(user_text or "").lower()).strip()
    if not text:
        return True
    word_count = len(text.split())
    if any(term in text for term in SOURCE_REQUEST_TERMS):
        return False
    try:
        risk = detect_high_risk_question(text, intent=intent)
        if risk.get("is_high_risk"):
            return False
    except Exception:
        pass
    if str(answer_mode or "") in {"riset", "kritis"}:
        return False
    if str(intent or "") == "quick_chat" and word_count <= 18:
        return True
    if text in CASUAL_CHAT_PATTERNS:
        return True
    if word_count <= 5 and any(text.startswith(prefix) for prefix in ["halo", "hai", "hi", "hello", "pagi", "siang", "sore", "malam", "makasih", "thanks"]):
        return True
    return False


def should_show_kb_sources_for_answer(
    user_text: str,
    intent: str = "",
    answer_mode: str = "auto",
    guard: Any = None,
    *,
    strict_rag_mode: bool = False,
    hide_for_casual: bool = True,
) -> bool:
    """Decide whether KB source details should be displayed to the user."""
    mode = normalize_answer_mode(answer_mode or "auto")
    lower = str(user_text or "").lower()
    if mode in {"riset", "kritis"}:
        return True
    if strict_rag_mode:
        return True
    try:
        if guard is not None and bool(getattr(guard, "is_high_risk", False)):
            return True
    except Exception:
        pass
    if any(term in lower for term in SOURCE_REQUEST_TERMS):
        return True
    if str(intent or "") in SOURCE_WORTHY_INTENTS:
        return True
    if hide_for_casual and is_casual_or_light_chat(user_text, intent=intent, answer_mode=mode):
        return False
    return False

def estimate_cost_idr(meta: Dict[str, Any], model: str = "") -> float:
    meta = meta or {}
    model_name = str(model or meta.get("active_model_final") or meta.get("model_requested") or meta.get("model") or "")
    usage = meta.get("usage") or {}
    in_tokens = int(usage.get("prompt_tokens") or usage.get("input_tokens") or 0)
    out_tokens = int(usage.get("completion_tokens") or usage.get("output_tokens") or 0)
    price = model_price(model_name)
    return round((in_tokens / 1_000_000) * int(price.get("input") or 0) + (out_tokens / 1_000_000) * int(price.get("output") or 0), 4)


def should_self_verify(intent: str, user_text: str, enabled: bool = False) -> bool:
    if not enabled:
        return False
    if intent in {"coding", "academic", "calculation", "deep_reasoning", "research", "document_question", "critical_current", "health", "livestock"}:
        return True
    t = str(user_text or "").lower()
    return any(x in t for x in ["pastikan", "cek lagi", "valid", "akurat", "jangan salah"])


def verify_answer_with_model(
    api_url: str,
    api_key: str,
    verifier_model: str,
    system_prompt: str,
    user_text: str,
    answer: str,
    temperature: float = 0.1,
    max_completion_tokens: int = 2200,
    timeout: int = 60,
) -> Tuple[str, Dict[str, Any]]:
    if not verifier_model:
        return answer, {"self_verification_skipped": "no_verifier_model"}
    messages = [
        {"role": "system", "content": (system_prompt or "Kamu adalah pemeriksa jawaban yang teliti.")[:2200]},
        {
            "role": "user",
            "content": (
                "Periksa jawaban berikut. Jika sudah benar, rapikan sedikit tanpa memperpanjang. "
                "Jika ada kekeliruan, perbaiki langsung. Jangan sebut proses verifikasi.\n\n"
                f"Pertanyaan pengguna:\n{str(user_text)[:4500]}\n\nJawaban awal:\n{str(answer)[:5500]}"
            ),
        },
    ]
    started = time.time()
    verified, meta = call_api_once(
        api_url=api_url,
        api_key=api_key,
        model=verifier_model,
        messages=messages,
        temperature=temperature,
        max_completion_tokens=max_completion_tokens,
        timeout=timeout,
    )
    meta["self_verified_by"] = verifier_model
    meta["self_verification_latency_seconds"] = round(time.time() - started, 3)
    return verified or answer, meta


# =========================
# Main answer wrapper
# =========================

def generate_power_answer(
    *,
    api_url: str,
    api_key: str,
    model: str,
    system_prompt: str,
    user_text: str,
    base_memory_text: str = "",
    recent_messages: Optional[List[Dict[str, str]]] = None,
    fallback_models: Optional[List[str]] = None,
    expensive_fallback_models: Optional[List[str]] = None,
    allow_expensive_fallback: bool = True,
    max_expensive_models: int = 1,
    temperature: float = 0.3,
    max_completion_tokens: int = 1800,
    timeout: int = 60,
    smart_model_router: bool = True,
    return_to_primary: bool = True,
    max_smart_models: int = 2,
    store: Optional[PowerStore] = None,
    user_id: str = "public",
    channel: str = "web",
    enable_rag: bool = True,
    rag_top_k: int = 5,
    enable_persistent_memory: bool = True,
    enable_prompt_templates: bool = True,
    enable_self_verification: bool = False,
    daily_cost_limit_idr: float = 0,
    max_expensive_calls_per_day: int = 0,
    enable_response_cache: bool = True,
    response_cache_ttl_seconds: int = 1800,
    enable_adaptive_scoring: bool = True,
    enable_circuit_breaker: bool = True,
    circuit_max_failures: int = 3,
    circuit_cooldown_seconds: int = 1800,
    anti_hallucination_enabled: bool = True,
    anti_hallucination_auto_strict: bool = True,
    anti_hallucination_min_sources: int = 1,
    anti_hallucination_min_quality: float = 0.0,
    anti_hallucination_min_freshness: float = 0.0,
    anti_hallucination_append_sources: bool = True,
    strict_rag_mode: bool = False,
    rag_min_sources: int = 1,
    rag_min_score: float = 0.0,
    quality_control_enabled: bool = True,
    quality_verifier_enabled: bool = True,
    quality_verifier_model: str = "",
    quality_min_score: float = 0.72,
    answer_mode: str = "auto",
    append_quality_footer: bool = False,
    hide_kb_sources_for_casual: bool = True,
    disable_rag_for_casual: bool = True,
    performance_optimizer_enabled: bool = True,
    query_rewriter_enabled: bool = True,
    reranker_enabled: bool = True,
    semantic_cache_enabled: bool = True,
    semantic_cache_threshold: float = 0.78,
    semantic_cache_ttl_seconds: int = 86400,
    latency_budget_enabled: bool = True,
    retrieval_eval_enabled: bool = True,
    live_music_chart_enabled: bool = True,
    live_music_chart_limit: int = 10,
    live_music_chart_timeout_seconds: int = 8,
    live_web_fallback_enabled: bool = True,
    live_web_fallback_provider: str = "tavily",
    tavily_api_key: str = "",
    live_web_fallback_max_results: int = 4,
    live_web_fallback_timeout_seconds: int = 10,
    live_web_fallback_min_sources: int = 1,
    live_web_fallback_include_raw_content: bool = True,
    live_web_fallback_max_content_chars: int = 3200,
    live_web_fallback_auto_save_to_kb: bool = True,
    live_web_fallback_ttl_hours: int = 24,
    live_web_fallback_force_for_current: bool = True,
    live_web_fallback_topic: str = "auto",
    **compat_kwargs: Any,
) -> Tuple[str, Dict[str, Any]]:
    store = store or get_power_store()
    ignored_compat_kwargs = sorted(str(key) for key in (compat_kwargs or {}).keys())
    intent = classify_intent_text(user_text)
    stored_mode = "auto"
    try:
        stored_mode = store.get_user_answer_mode(user_id, default="auto")
    except Exception:
        stored_mode = "auto"
    requested_mode = normalize_answer_mode(answer_mode or stored_mode or "auto")
    if requested_mode == "auto" and stored_mode != "auto":
        requested_mode = normalize_answer_mode(stored_mode)
    effective_answer_mode = infer_answer_mode(user_text, requested_mode=requested_mode, intent=intent)
    answer_mode_policy = mode_policy(effective_answer_mode)
    if bool(answer_mode_policy.get("force_rag")):
        enable_rag = True
    if bool(answer_mode_policy.get("strict_rag")):
        strict_rag_mode = True
        anti_hallucination_auto_strict = True
    if int(answer_mode_policy.get("min_sources") or 0) > 0:
        anti_hallucination_min_sources = max(int(anti_hallucination_min_sources or 1), int(answer_mode_policy.get("min_sources") or 1))
        rag_min_sources = max(int(rag_min_sources or 1), int(answer_mode_policy.get("min_sources") or 1))
    try:
        temperature = min(float(temperature or 0.3), float(answer_mode_policy.get("temperature_cap") or 0.3))
    except Exception:
        temperature = 0.3

    music_chart_query = bool(is_music_chart_query(user_text))
    # Tangga lagu terbaru adalah data aktual, tetapi risikonya rendah.
    # Jangan biarkan guard mengubah pertanyaan hiburan ringan menjadi blokir total
    # selama mode pengguna bukan riset/kritis/strict. Bukti tetap diambil dari helper live chart.
    if music_chart_query and effective_answer_mode not in {"riset", "kritis"}:
        anti_hallucination_auto_strict = False
        strict_rag_mode = False
        anti_hallucination_min_sources = min(int(anti_hallucination_min_sources or 1), 1)
        rag_min_sources = min(int(rag_min_sources or 1), 1)
        # Chart musik berubah cepat; hindari cache jawaban lama.
        enable_response_cache = False
        semantic_cache_enabled = False

    query_plan = rewrite_query(user_text, intent=intent, answer_mode=effective_answer_mode) if (performance_optimizer_enabled and query_rewriter_enabled) else QueryPlan(user_text, user_text, [], False, True, "disabled")
    retrieval_query = str(getattr(query_plan, "rewritten_query", "") or user_text)
    if performance_optimizer_enabled and latency_budget_enabled:
        try:
            timeout = max(4, min(int(timeout or 60), int(latency_budget_seconds(effective_answer_mode, intent, user_text))))
        except Exception:
            pass

    casual_rag_skipped = False
    if (
        enable_rag
        and bool(disable_rag_for_casual)
        and not bool(answer_mode_policy.get("force_rag"))
        and is_casual_or_light_chat(user_text, intent=intent, answer_mode=effective_answer_mode)
    ):
        enable_rag = False
        casual_rag_skipped = True

    if daily_cost_limit_idr and daily_cost_limit_idr > 0:
        usage = store.usage_summary(days=1)
        if float(usage.get("cost_idr") or 0) >= float(daily_cost_limit_idr):
            return (
                "Batas biaya harian AI sudah tercapai. Admin dapat menaikkan DAILY_COST_LIMIT_IDR atau menunggu reset hari berikutnya.",
                {"budget_guard_blocked": True, "intent": intent, "daily_cost_limit_idr": daily_cost_limit_idr},
            )
    if max_expensive_calls_per_day and max_expensive_calls_per_day > 0:
        if store.count_expensive_calls_today() >= int(max_expensive_calls_per_day):
            allow_expensive_fallback = False
            expensive_fallback_models = []

    rag_sources: List[Dict[str, Any]] = []
    retrieval_started = time.time()
    if enable_rag:
        try:
            raw_limit = max(1, int(rag_top_k or 5))
            search_limit = raw_limit * 3 if (performance_optimizer_enabled and reranker_enabled) else raw_limit
            rag_sources = store.search_documents(retrieval_query, limit=search_limit, min_score=float(rag_min_score or 0.0))
            if performance_optimizer_enabled and reranker_enabled:
                rag_sources = rerank_sources(retrieval_query, rag_sources, limit=raw_limit, diversity=True)
        except Exception:
            rag_sources = []

    music_chart_result = None
    if music_chart_query and bool(live_music_chart_enabled):
        try:
            music_chart_result = fetch_indonesia_music_charts(
                limit=int(live_music_chart_limit or 10),
                timeout=int(live_music_chart_timeout_seconds or 8),
            )
            if getattr(music_chart_result, "ok", False):
                pseudo = music_chart_result_to_pseudo_source(music_chart_result)
                if pseudo:
                    rag_sources = [pseudo] + list(rag_sources or [])
                    enable_rag = True
            elif not rag_sources:
                answer = build_music_chart_fallback_answer(music_chart_result)
                meta = {
                    "intent": intent,
                    "music_chart_query": True,
                    "music_chart_live_fetch_ok": False,
                    "music_chart_result": music_chart_result.to_dict() if hasattr(music_chart_result, "to_dict") else {},
                    "anti_hallucination_blocked": False,
                    "show_kb_sources": False,
                    "kb_source_display_policy": "music_chart_live_fetch_failed",
                }
                try:
                    store.log_interaction(user_id=user_id, channel=channel, intent=intent, model=model, question=user_text, answer=answer, meta=meta, success=True)
                except Exception:
                    pass
                return answer, meta
        except Exception as exc:
            if not rag_sources:
                answer = (
                    "Saya belum bisa mengambil tangga lagu Indonesia terbaru saat ini. "
                    "Coba jalankan /update atau tambahkan sumber chart musik seperti Billboard Indonesia Songs, Spotify Top 50 Indonesia, Apple Music Top Charts Indonesia, dan YouTube Charts Indonesia."
                )
                meta = {"intent": intent, "music_chart_query": True, "music_chart_live_fetch_error": str(exc)[:500], "show_kb_sources": False}
                return answer, meta


    live_search_result = None
    live_search_decision: Dict[str, Any] = {"use": False, "reason": "disabled"}
    live_search_save_meta: Dict[str, Any] = {}
    if bool(live_web_fallback_enabled) and str(live_web_fallback_provider or "tavily").lower() == "tavily":
        try:
            live_min_sources = max(1, int(live_web_fallback_min_sources or 1))
            live_search_decision = should_trigger_live_fallback(
                user_text,
                intent=intent,
                answer_mode=effective_answer_mode,
                rag_sources=rag_sources,
                min_sources=live_min_sources,
                force=False,
            )
            # For current factual questions, live search can supplement stale KB even outside strict mode.
            if bool(live_web_fallback_force_for_current) and not live_search_decision.get("use"):
                live_search_decision = should_trigger_live_fallback(
                    user_text,
                    intent=intent,
                    answer_mode=effective_answer_mode,
                    rag_sources=rag_sources,
                    min_sources=live_min_sources,
                    force=False,
                )
            if bool(live_search_decision.get("use")):
                live_search_result = tavily_live_search(
                    retrieval_query,
                    api_key=str(tavily_api_key or ""),
                    max_results=int(live_web_fallback_max_results or 4),
                    timeout=int(live_web_fallback_timeout_seconds or 10),
                    include_raw_content=bool(live_web_fallback_include_raw_content),
                    max_content_chars=int(live_web_fallback_max_content_chars or 3200),
                    topic=str(live_web_fallback_topic or "auto"),
                )
                if getattr(live_search_result, "ok", False):
                    live_sources = live_result_to_rag_sources(
                        live_search_result,
                        max_sources=int(live_web_fallback_max_results or 4),
                    )
                    rag_sources = list(live_sources or []) + list(rag_sources or [])
                    # Prefer fresh live sources, but keep a few KB sources for grounding diversity.
                    if performance_optimizer_enabled and reranker_enabled:
                        rag_sources = rerank_sources(
                            retrieval_query,
                            rag_sources,
                            limit=max(int(rag_top_k or 5), int(live_web_fallback_max_results or 4)),
                            diversity=True,
                        )
                    enable_rag = True
                    # Current/live answers should not be cached too aggressively.
                    enable_response_cache = False
                    semantic_cache_enabled = False
                    if bool(live_web_fallback_auto_save_to_kb):
                        live_search_save_meta = save_live_result_to_kb(
                            store,
                            live_search_result,
                            ttl_hours=int(live_web_fallback_ttl_hours or 24),
                            max_items=int(live_web_fallback_max_results or 4),
                        )
        except Exception as exc:
            live_search_decision = {"use": False, "reason": "error", "error": str(exc)[:500]}

    retrieval_latency = round(time.time() - retrieval_started, 4)
    retrieval_metrics = retrieval_precision_estimate(retrieval_query, rag_sources) if (performance_optimizer_enabled and retrieval_eval_enabled) else {}
    if performance_optimizer_enabled and retrieval_eval_enabled and enable_rag:
        try:
            store.record_retrieval_eval(
                question=user_text,
                search_query=retrieval_query,
                sources=rag_sources,
                intent=intent,
                user_id=user_id,
                channel=channel,
                latency_seconds=retrieval_latency,
                meta={"query_plan": query_plan.to_dict() if hasattr(query_plan, "to_dict") else {}, "metrics": retrieval_metrics},
            )
        except Exception:
            pass
    guard_result = evaluate_evidence_gate(
        user_text,
        rag_sources,
        intent=intent,
        enabled=bool(anti_hallucination_enabled),
        auto_strict=bool(anti_hallucination_auto_strict),
        strict_rag_mode=bool(strict_rag_mode),
        min_sources=max(int(rag_min_sources or 1), int(anti_hallucination_min_sources or 1)),
        min_quality=float(anti_hallucination_min_quality or 0.0),
        min_freshness=float(anti_hallucination_min_freshness or 0.0),
    )
    show_kb_sources = should_show_kb_sources_for_answer(
        user_text,
        intent=intent,
        answer_mode=effective_answer_mode,
        guard=guard_result,
        strict_rag_mode=bool(strict_rag_mode),
        hide_for_casual=bool(hide_kb_sources_for_casual),
    )

    if enable_rag and getattr(guard_result, "enabled", False) and not getattr(guard_result, "allow_answer", True):
        gap_id = store.log_knowledge_gap(
            question=user_text,
            reason=str(getattr(guard_result, "reason", "insufficient_evidence")),
            intent=intent,
            user_id=user_id,
            channel=channel,
            priority=1,
            suggested_query=user_text,
            meta={
                "guard": guard_result.to_dict() if hasattr(guard_result, "to_dict") else {},
                "rag_sources_found": len(rag_sources),
            },
        )
        return (
            build_insufficient_evidence_answer(user_text, guard_result, intent=intent),
            {
                "anti_hallucination_blocked": True,
                "strict_rag_blocked": bool(strict_rag_mode),
                "knowledge_gap_id": gap_id,
                "intent": intent,
                "power_kb_sources": rag_sources,
                "show_kb_sources": True,
                "kb_source_display_policy": "forced_insufficient_evidence",
                "hallucination_guard": guard_result.to_dict() if hasattr(guard_result, "to_dict") else {},
            },
        )

    memory_text = build_power_context(
        store=store,
        user_text=user_text,
        base_memory=base_memory_text,
        user_id=user_id,
        enable_rag=enable_rag,
        rag_top_k=max(1, int(rag_top_k or 5)),
        enable_persistent_memory=enable_persistent_memory,
        retrieval_query=retrieval_query,
        preselected_docs=rag_sources,
    )
    routed_user_text = enhance_prompt_for_intent(user_text, intent, enable_templates=enable_prompt_templates)
    guard_instruction = build_guard_system_instruction(user_text, rag_sources, guard_result) if anti_hallucination_enabled else ""
    mode_instruction = build_mode_system_instruction(effective_answer_mode, user_text) if quality_control_enabled else ""
    live_instruction = build_live_search_system_instruction(live_search_result) if live_search_result is not None else ""
    guarded_system_prompt = (
        str(system_prompt or "").strip()
        + ("\n\n" + mode_instruction if mode_instruction else "")
        + ("\n\n" + guard_instruction if guard_instruction else "")
        + ("\n\n" + live_instruction if live_instruction else "")
    ).strip()
    guarded_temperature = apply_temperature_policy(float(temperature or 0.3), guard_result)

    # Adaptive policy: rank candidate models using historical success, quality, latency, cost, and circuit breaker.
    cheap_candidates = list(dict.fromkeys([m for m in (fallback_models or []) if m]))
    expensive_candidates = list(dict.fromkeys([m for m in (expensive_fallback_models or []) if m]))
    all_candidates = list(dict.fromkeys([model] + cheap_candidates + expensive_candidates))
    if enable_circuit_breaker:
        all_candidates = store.filter_blocked_models(all_candidates)
    if enable_adaptive_scoring:
        ranked_all = store.rank_models_for_intent(all_candidates, intent)
    else:
        ranked_all = all_candidates
    selected_model = ranked_all[0] if ranked_all else model

    ranked_cheap = [m for m in ranked_all if m != selected_model and model_cost_tier(m) == "cheap"]
    ranked_expensive = [m for m in ranked_all if m != selected_model and model_cost_tier(m) in {"medium", "expensive", "ultra"}]
    # Preserve any still unranked fallback if all candidates were filtered weirdly.
    ranked_cheap.extend([m for m in cheap_candidates if m != selected_model and m not in ranked_cheap])
    ranked_expensive.extend([m for m in expensive_candidates if m != selected_model and m not in ranked_expensive])

    adjusted_max_tokens = adaptive_token_budget_for_intent(intent, user_text, base=max_completion_tokens)
    try:
        adjusted_max_tokens = int(max(500, min(6000, adjusted_max_tokens * float(answer_mode_policy.get("token_multiplier") or 1.0))))
    except Exception:
        pass
    route_signature = ",".join(ranked_all[:8]) + f"|show_kb_sources={int(bool(show_kb_sources))}|casual_rag_skipped={int(bool(casual_rag_skipped))}|retrieval_query={hashlib.sha256(str(retrieval_query).encode('utf-8')).hexdigest()[:12]}"
    cache_key = make_response_cache_key(
        model=selected_model,
        system_prompt=system_prompt,
        user_text=user_text,
        memory_text=memory_text,
        intent=intent,
        route_signature=route_signature,
    )
    if enable_response_cache and intent not in {"admin_command", "coding"}:
        cached = store.get_cached_response(cache_key)
        if cached:
            answer, meta = cached
            meta["power_intent"] = intent
            meta["active_model_final"] = meta.get("active_model_final") or selected_model
            meta.setdefault("show_kb_sources", bool(show_kb_sources))
            meta.setdefault("kb_source_display_policy", "auto")
            meta.setdefault("casual_rag_skipped", bool(casual_rag_skipped))
            meta.setdefault("query_plan", query_plan.to_dict() if hasattr(query_plan, "to_dict") else {})
            return answer, meta
        if (
            performance_optimizer_enabled
            and semantic_cache_enabled
            and not bool(strict_rag_mode)
            and effective_answer_mode not in {"riset", "kritis"}
            and not getattr(guard_result, "strict", False)
        ):
            semantic_cached = store.get_semantic_cached_response(
                user_text,
                intent=intent,
                threshold=float(semantic_cache_threshold or 0.78),
                ttl_seconds=int(semantic_cache_ttl_seconds or 86400),
                user_id=user_id,
                channel=channel,
            )
            if semantic_cached:
                answer, meta = semantic_cached
                meta["power_intent"] = intent
                meta["active_model_final"] = meta.get("active_model_final") or selected_model
                meta.setdefault("show_kb_sources", False if casual_rag_skipped else bool(show_kb_sources))
                meta.setdefault("kb_source_display_policy", "semantic_cache")
                meta.setdefault("casual_rag_skipped", bool(casual_rag_skipped))
                meta.setdefault("query_plan", query_plan.to_dict() if hasattr(query_plan, "to_dict") else {})
                return answer, meta

    started = time.time()
    answer = ""
    meta: Dict[str, Any] = {}
    success = False
    final_model = selected_model
    last_error = ""

    def _call_once(primary: str, cheap_pool: List[str], expensive_pool: List[str]) -> Tuple[str, Dict[str, Any]]:
        return __import__("ai_core").generate_answer(
            api_url=api_url,
            api_key=api_key,
            model=primary,
            system_prompt=guarded_system_prompt,
            user_text=routed_user_text,
            memory_text=memory_text,
            recent_messages=recent_messages or [],
            fallback_models=cheap_pool,
            expensive_fallback_models=expensive_pool,
            allow_expensive_fallback=allow_expensive_fallback and bool(expensive_pool),
            max_expensive_models=max_expensive_models,
            temperature=guarded_temperature,
            max_completion_tokens=adjusted_max_tokens,
            timeout=timeout,
            smart_model_router=smart_model_router,
            return_to_primary=return_to_primary,
            max_smart_models=max_smart_models,
        )

    try:
        try:
            answer, meta = _call_once(selected_model, ranked_cheap, ranked_expensive)
        except Exception as exc:
            last_error = str(exc)
            if enable_circuit_breaker:
                store.register_model_failure(selected_model, error=last_error, max_failures=circuit_max_failures, cooldown_seconds=circuit_cooldown_seconds)
            # Retry once with the next ranked model, if available.
            alternates = [m for m in ranked_all if m != selected_model]
            if not alternates:
                raise
            retry_model = alternates[0]
            retry_cheap = [m for m in ranked_cheap if m != retry_model]
            retry_expensive = [m for m in ranked_expensive if m != retry_model]
            answer, meta = _call_once(retry_model, retry_cheap, retry_expensive)
            meta = meta or {}
            meta["power_auto_retry_from"] = selected_model
            meta["power_auto_retry_to"] = retry_model
            meta["power_auto_retry_error"] = last_error[:800]
            selected_model = retry_model

        meta = meta or {}
        final_model = str(meta.get("active_model_final") or meta.get("model_requested") or selected_model)
        meta["power_intent"] = intent
        meta["answer_mode"] = effective_answer_mode
        meta["answer_mode_requested"] = requested_mode
        meta["answer_mode_policy"] = answer_mode_policy
        meta["power_rag_enabled"] = bool(enable_rag)
        meta["show_kb_sources"] = bool(show_kb_sources)
        meta["kb_source_display_policy"] = "auto"
        meta["casual_rag_skipped"] = bool(casual_rag_skipped)
        kb_source_items = [
            {
                "doc_id": item.get("doc_id"),
                "title": item.get("title"),
                "source": item.get("source"),
                "collection": item.get("collection"),
                "tags": item.get("tags"),
                "heading": item.get("heading"),
                "page_label": item.get("page_label"),
                "chunk_index": item.get("chunk_index"),
                "score": item.get("score"),
                "source_quality": item.get("source_quality"),
                "freshness_score": item.get("freshness_score"),
                "published_at": item.get("published_at"),
                "criticality_score": item.get("criticality_score"),
                "citation": item.get("citation") or _format_kb_citation(item),
            }
            for item in (rag_sources or [])[:5]
        ]
        meta["power_kb_sources"] = kb_source_items
        meta["power_rag_sources"] = kb_source_items
        meta["power_persistent_memory_enabled"] = bool(enable_persistent_memory)
        meta["power_prompt_template_enabled"] = bool(enable_prompt_templates)
        meta["power_adaptive_scoring_enabled"] = bool(enable_adaptive_scoring)
        meta["power_circuit_breaker_enabled"] = bool(enable_circuit_breaker)
        meta["power_ranked_models"] = ranked_all[:10]
        meta["power_selected_model"] = selected_model
        meta["power_adjusted_max_tokens"] = adjusted_max_tokens
        meta["power_latency_seconds"] = round(time.time() - started, 3)
        meta["query_plan"] = query_plan.to_dict() if hasattr(query_plan, "to_dict") else {}
        meta["retrieval_query"] = retrieval_query
        meta["retrieval_latency_seconds"] = retrieval_latency
        meta["retrieval_metrics"] = retrieval_metrics
        meta["performance_optimizer_enabled"] = bool(performance_optimizer_enabled)
        meta["live_web_fallback_enabled"] = bool(live_web_fallback_enabled)
        meta["live_web_fallback_provider"] = str(live_web_fallback_provider or "tavily")
        meta["live_web_fallback_decision"] = live_search_decision
        if live_search_result is not None and hasattr(live_search_result, "to_dict"):
            meta["live_web_fallback_result"] = live_search_result.to_dict()
            meta["live_web_fallback_used"] = bool(getattr(live_search_result, "ok", False))
        else:
            meta["live_web_fallback_used"] = False
        if live_search_save_meta:
            meta["live_web_fallback_saved_to_kb"] = live_search_save_meta
        if music_chart_query:
            meta["music_chart_query"] = True
            meta["music_chart_live_fetch_ok"] = bool(getattr(music_chart_result, "ok", False)) if music_chart_result is not None else False
            if music_chart_result is not None and hasattr(music_chart_result, "to_dict"):
                meta["music_chart_result"] = music_chart_result.to_dict()
        if ignored_compat_kwargs:
            meta["ignored_compat_kwargs"] = ignored_compat_kwargs

        if should_self_verify(intent, user_text, enabled=enable_self_verification):
            verifier = ""
            if ranked_expensive:
                verifier = ranked_expensive[0]
            elif ranked_cheap:
                verifier = ranked_cheap[0]
            else:
                verifier = final_model
            try:
                verified_answer, verify_meta = verify_answer_with_model(
                    api_url=api_url,
                    api_key=api_key,
                    verifier_model=verifier,
                    system_prompt=guarded_system_prompt,
                    user_text=user_text,
                    answer=answer,
                    temperature=min(float(guarded_temperature), 0.2),
                    max_completion_tokens=max(adjusted_max_tokens, 2200),
                    timeout=timeout,
                )
                if verified_answer:
                    answer = verified_answer
                    meta["self_verification"] = verify_meta
                    meta["self_verified_by"] = verifier
            except Exception as exc:
                meta["self_verification_error"] = str(exc)[:500]
        if anti_hallucination_enabled:
            try:
                answer = append_guard_note(
                    answer,
                    rag_sources,
                    guard=guard_result,
                    append_sources=bool(anti_hallucination_append_sources) and bool(show_kb_sources),
                )
                meta["hallucination_guard"] = guard_result.to_dict() if hasattr(guard_result, "to_dict") else {}
                meta["hallucination_guard_temperature"] = guarded_temperature
                meta["hallucination_claim_risk"] = lightweight_claim_risk(answer)
            except Exception as exc:
                meta["hallucination_guard_error"] = str(exc)[:500]
        if quality_control_enabled:
            try:
                quality_result = score_answer_quality(
                    question=user_text,
                    answer=answer,
                    rag_sources=rag_sources,
                    mode=effective_answer_mode,
                    intent=intent,
                    guard_meta=guard_result.to_dict() if hasattr(guard_result, "to_dict") else {},
                )
                meta["answer_quality"] = quality_result.to_dict() if hasattr(quality_result, "to_dict") else {}
                should_verify_quality = bool(quality_verifier_enabled) and (
                    bool(answer_mode_policy.get("verifier")) or bool(getattr(quality_result, "needs_verification", False))
                ) and float(getattr(quality_result, "score", 1.0) or 1.0) < float(quality_min_score or 0.72)
                if should_verify_quality:
                    verifier = str(quality_verifier_model or "").strip()
                    if not verifier:
                        if ranked_expensive:
                            verifier = ranked_expensive[0]
                        elif ranked_cheap:
                            verifier = ranked_cheap[0]
                        else:
                            verifier = final_model
                    repaired_answer, verifier_meta = verify_and_repair_answer(
                        api_url=api_url,
                        api_key=api_key,
                        verifier_model=verifier,
                        system_prompt=guarded_system_prompt,
                        question=user_text,
                        answer=answer,
                        rag_sources=rag_sources,
                        mode=effective_answer_mode,
                        timeout=timeout,
                        max_completion_tokens=max(adjusted_max_tokens, 2200),
                    )
                    if repaired_answer and repaired_answer.strip():
                        answer = repaired_answer.strip()
                    meta["quality_verification"] = verifier_meta
                    meta["quality_verified_by"] = verifier
                    quality_result = score_answer_quality(
                        question=user_text,
                        answer=answer,
                        rag_sources=rag_sources,
                        mode=effective_answer_mode,
                        intent=intent,
                        guard_meta=guard_result.to_dict() if hasattr(guard_result, "to_dict") else {},
                    )
                    meta["answer_quality_after_verifier"] = quality_result.to_dict() if hasattr(quality_result, "to_dict") else {}
                if append_quality_footer:
                    answer = answer + build_quality_footer(quality_result, effective_answer_mode)
            except Exception as exc:
                meta["quality_control_error"] = str(exc)[:500]

        success = True
        if enable_response_cache and intent not in {"admin_command", "coding"}:
            store.set_cached_response(cache_key, answer, meta, ttl_seconds=response_cache_ttl_seconds)
            if (
                performance_optimizer_enabled
                and semantic_cache_enabled
                and not bool(strict_rag_mode)
                and effective_answer_mode not in {"riset", "kritis"}
                and len(str(answer or "")) >= 30
            ):
                try:
                    store.set_semantic_cached_response(
                        question=user_text,
                        answer=answer,
                        meta=meta,
                        intent=intent,
                        ttl_seconds=int(semantic_cache_ttl_seconds or 86400),
                        user_id=user_id,
                        channel=channel,
                    )
                except Exception:
                    pass
        return answer, meta
    finally:
        try:
            if not success and last_error:
                meta["error"] = last_error[:1000]
            final_model = str((meta or {}).get("active_model_final") or (meta or {}).get("model_requested") or final_model or selected_model or model)
            interaction_id = store.log_interaction(
                user_id=user_id,
                channel=channel,
                intent=intent,
                model=final_model,
                question=user_text,
                answer=answer,
                meta=meta,
                success=success,
            )
            if isinstance(meta, dict):
                meta["power_interaction_id"] = interaction_id
            try:
                qd = (meta or {}).get("answer_quality_after_verifier") or (meta or {}).get("answer_quality") or {}
                if qd:
                    qr_id = store.record_answer_quality(
                        interaction_id=interaction_id,
                        user_id=user_id,
                        channel=channel,
                        intent=intent,
                        answer_mode=effective_answer_mode,
                        score=float(qd.get("score") or 0),
                        level=str(qd.get("level") or ""),
                        needs_verification=bool(qd.get("needs_verification")),
                        verifier_model=str((meta or {}).get("quality_verified_by") or ""),
                        verified=bool((meta or {}).get("quality_verification")),
                        reasons=list(qd.get("reasons") or []),
                        metrics=dict(qd.get("metrics") or {}),
                        question=user_text,
                        answer=answer,
                    )
                    if isinstance(meta, dict):
                        meta["answer_quality_report_id"] = qr_id
            except Exception:
                pass
            if success and enable_rag and not rag_sources and intent in {"research", "academic", "document_question", "health", "livestock"}:
                meta["knowledge_gap_id"] = store.log_knowledge_gap(
                    question=user_text,
                    reason="no_relevant_kb_source",
                    intent=intent,
                    user_id=user_id,
                    channel=channel,
                    priority=2,
                    suggested_query=user_text,
                )
        except Exception:
            pass


# =========================
# Commands / benchmark
# =========================

def handle_power_command(text: str, store: PowerStore, user_id: str = "global", is_admin: bool = False) -> str:
    raw = str(text or "").strip()
    lower = raw.lower()
    if not raw.startswith("/"):
        return ""

    admin_only_prefixes = (
        "/ingat", "/lupa", "/rag", "/kb", "/biaya", "/usage", "/dokumen", "/benchmark",
        "/model skor", "/model score", "/circuit", "/cache bersih", "/cache clear",
        "/briefing", "/trending", "/cek isu", "/pantau", "/kualitas", "/quality", "/evaluasi", "/laporan", "/export", "/performa", "/performance", "/optimasi",
    )
    if lower.startswith(admin_only_prefixes) and not is_admin:
        return "Perintah ini hanya untuk admin."

    mode_cmd = parse_mode_command(raw)
    if mode_cmd:
        current = store.get_user_answer_mode(user_id, default="auto")
        if mode_cmd == "list":
            return mode_help_text(current)
        selected = store.set_user_answer_mode(user_id, mode_cmd)
        return "✅ Mode jawaban diubah menjadi: " + selected + "\n\n" + mode_help_text(selected)

    if lower in {"/performa", "/performance", "/performa ai", "/performance ai"}:
        data = store.performance_dashboard(days=14)
        retrieval = data.get("retrieval") or {}
        lines = [
            "⚡ Performance AI",
            "",
            f"Evaluasi retrieval: {retrieval.get('total', 0)}",
            f"Precision rata-rata: {float(retrieval.get('avg_precision') or 0):.2f}",
            f"Similarity rata-rata: {float(retrieval.get('avg_similarity') or 0):.2f}",
            f"Latency retrieval: {float(retrieval.get('avg_retrieval_latency') or 0):.3f}s",
            f"Semantic cache aktif: {data.get('semantic_cache_active', 0)}",
            f"Semantic cache hit 14 hari: {data.get('semantic_cache_hits', 0)}",
            "",
            "Intent paling banyak / latency:",
        ]
        for row in data.get("top_intents_latency", [])[:8]:
            lines.append(f"- {row.get('intent') or 'unknown'}: {row.get('jumlah')} | avg {row.get('avg_latency')}s")
        return "\n".join(lines)

    if lower in {"/optimasi db", "/optimize db", "/maintenance db"}:
        result = store.optimize_database(vacuum=False)
        return "✅ Optimasi DB selesai.\n" + json.dumps(result, ensure_ascii=False, indent=2)[:1800]

    if lower in {"/kualitas", "/quality", "/kualitas ai", "/quality ai"}:
        data = store.quality_dashboard(days=14)
        lines = [
            "📊 Quality Control AI",
            "",
            f"Total dinilai: {data.get('total', 0)}",
            f"Rata-rata skor: {data.get('avg_score', 0)}",
            f"Skor rendah: {data.get('low_count', 0)}",
            f"Diverifikasi: {data.get('verified_count', 0)}",
            "",
            "Per mode:",
        ]
        for row in data.get("by_mode", [])[:8]:
            lines.append(f"- {row.get('answer_mode')}: {row.get('jumlah')} | avg {float(row.get('avg_score') or 0):.2f}")
        low_items = data.get("low_quality", [])[:5]
        if low_items:
            lines.append("")
            lines.append("Jawaban skor rendah terbaru:")
            for item in low_items:
                lines.append(f"- #{item.get('id')} skor {float(item.get('score') or 0):.2f}: {str(item.get('question_preview') or '')[:100]}")
        return "\n".join(lines)

    if lower in {"/evaluasi mingguan", "/laporan mingguan", "/weekly report"}:
        return store.weekly_quality_evaluation(days=7, save=True)

    if lower in {"/export kb", "/export knowledge", "/export knowledgebase"}:
        return "Export KB tersedia di dashboard Admin web: tab Quality Control → Export/Import. Untuk Telegram, gunakan web agar file JSONL bisa diunduh dengan aman."

    if lower in {"/briefing", "/briefing harian", "/trending"}:
        return store.daily_current_briefing(days=1, limit=12)

    if lower.startswith("/cek isu ") or lower.startswith("/cekisu "):
        if lower.startswith("/cek isu "):
            query = raw.split(" ", 2)[2].strip()
        else:
            query = raw.split(" ", 1)[1].strip() if " " in raw else ""
        if not query:
            return "Format: /cek isu <topik>"
        return store.build_current_issue_report(query, limit=8)

    if lower in {"/pantau", "/pantau list", "/watchlist"}:
        topics = store.list_watch_topics(active_only=True, limit=50)
        if not topics:
            return "Watchlist masih kosong. Tambahkan dengan: /pantau <topik>"
        lines = ["👁️ Watchlist Isu", ""]
        for item in topics:
            lines.append(f"#{item.get('id')} [{item.get('priority')}] {item.get('topic')} — {item.get('category')}")
        return "\n".join(lines)

    if lower.startswith("/pantau hapus "):
        parts = raw.split()
        try:
            watch_id = int(parts[-1])
        except Exception:
            return "Format: /pantau hapus <id>"
        ok = store.remove_watch_topic(watch_id)
        return "✅ Watchlist dinonaktifkan." if ok else "ID watchlist tidak ditemukan."

    if lower.startswith("/pantau "):
        topic = raw.split(" ", 1)[1].strip()
        if not topic:
            return "Format: /pantau <topik>"
        watch_id = store.add_watch_topic(topic, category="manual", priority=2, meta={"created_by": user_id})
        return f"✅ Topik dipantau. ID: {watch_id}\nTopik: {topic}"

    if lower.startswith("/ingat "):
        body = raw.split(" ", 1)[1].strip()
        mem_id = store.add_memory(body, user_id=user_id)
        return f"✅ Memory permanen disimpan. ID: {mem_id}"

    if lower.startswith("/lupa "):
        key = raw.split(" ", 1)[1].strip()
        count = store.delete_memories_containing(key, user_id=None)
        return f"✅ Memory yang mengandung '{key}' dihapus: {count}."

    if lower in {"/kb", "/rag", "/kb bantuan", "/rag bantuan", "/kb help", "/rag help"}:
        return (
            "📚 Knowledge Base\n\n"
            "• /kb statistik — jumlah dokumen dan chunk\n"
            "• /kb list — daftar dokumen terakhir\n"
            "• /kb cari <query> — cari isi knowledge base\n"
            "• /kb detail <doc_id> — preview dokumen\n"
            "• /kb hapus <doc_id> — hapus dokumen\n"
            "• /kb rebuild — bangun ulang index FTS\n"
            "• /kb koleksi — daftar koleksi/workspace KB\n"
            "• /kb pin <doc_id> — prioritaskan dokumen\n"
            "• /kb unpin <doc_id> — lepas prioritas dokumen\n"
            "• /kb set <doc_id> <koleksi> | <tag1,tag2> — ubah metadata\n"
            "• /kb tambah <judul> lalu baris baru isi dokumen\n"
            "\n🧭 Critical Current Layer:\n"
            "• /briefing — ringkasan klaim/fakta terkini dari KB\n"
            "• /trending — sama dengan /briefing ringkas\n"
            "• /cek isu <topik> — cek isu kritis dengan klaim + dokumen pendukung\n"
            "• /pantau <topik> — tambah topik watchlist\n"
            "• /pantau list — daftar watchlist\n"
            "• /pantau hapus <id> — nonaktifkan watchlist"
        )

    if lower in {"/kb statistik", "/rag statistik", "/kb stats", "/rag stats"}:
        stats = store.knowledge_stats()
        return (
            "📚 Statistik Knowledge Base:\n"
            f"Dokumen: {stats.get('documents', 0)}\n"
            f"Chunks: {stats.get('chunks', 0)}\n"
            f"Karakter: {stats.get('characters', 0)}\n"
            f"Terakhir: {stats.get('latest_title', '-')} {stats.get('latest_at_wib', '')}"
        )

    if lower in {"/kb koleksi", "/rag koleksi", "/kb collections", "/rag collections"}:
        cols = store.knowledge_collections()
        if not cols:
            return "Knowledge base masih kosong."
        lines = ["📚 Koleksi Knowledge Base", ""]
        for item in cols[:30]:
            lines.append(f"• {item.get('collection')}: {item.get('documents')} dokumen | {item.get('chunks')} chunk")
        return "\n".join(lines)

    if lower.startswith(("/kb pin ", "/rag pin ", "/kb unpin ", "/rag unpin ")):
        parts = raw.split()
        if len(parts) < 3 or not parts[2].isdigit():
            return "Format: /kb pin <doc_id> atau /kb unpin <doc_id>"
        want_pin = parts[1].lower() == "pin"
        ok = store.set_document_pinned(int(parts[2]), pinned=want_pin)
        return ("✅ Dokumen diprioritaskan." if want_pin else "✅ Prioritas dokumen dilepas.") if ok else "Dokumen tidak ditemukan."

    if lower.startswith(("/kb set ", "/rag set ")):
        parts = raw.split(" ", 3)
        if len(parts) < 4 or not parts[2].isdigit():
            return "Format: /kb set <doc_id> <koleksi> | <tag1,tag2>"
        doc_id = int(parts[2])
        body = parts[3]
        if "|" in body:
            collection, tags = body.split("|", 1)
        else:
            collection, tags = body, ""
        ok = store.update_document_metadata(doc_id, collection=collection.strip(), tags=tags.strip())
        return "✅ Metadata dokumen diperbarui." if ok else "Dokumen tidak ditemukan."

    if lower in {"/kb list", "/rag list", "/kb daftar", "/rag daftar"}:
        docs = store.list_documents(limit=15)
        if not docs:
            return "Knowledge base masih kosong."
        lines = ["📚 Dokumen Knowledge Base terakhir", ""]
        for doc in docs:
            lines.append(
                f"ID {doc.get('id')} — {doc.get('title')}\n"
                f"Chunk: {doc.get('chunks')} | Tanggal: {doc.get('created_at_wib', '-')}"
            )
            lines.append("")
        return "\n".join(lines).strip()

    if lower.startswith(("/kb detail ", "/rag detail ")):
        doc_id = raw.split(" ", 2)[2].strip()
        doc = store.get_document(int(doc_id), max_chars=2200) if doc_id.isdigit() else {}
        if not doc:
            return "Dokumen tidak ditemukan."
        preview = re.sub(r"\s+", " ", str(doc.get("preview") or "")).strip()[:1800]
        return f"📄 {doc.get('title')}\nID: {doc.get('id')} | Source: {doc.get('source')} | Chunks: {doc.get('chunks')}\n\n{preview}"

    if lower.startswith(("/kb hapus ", "/rag hapus ", "/kb delete ", "/rag delete ")):
        doc_id = raw.split(" ", 2)[2].strip()
        ok = store.delete_document(int(doc_id)) if doc_id.isdigit() else False
        return f"✅ Dokumen ID {doc_id} dihapus." if ok else "Dokumen tidak ditemukan/gagal dihapus."

    if lower in {"/kb rebuild", "/rag rebuild", "/kb index", "/rag index"}:
        docs, chunks = store.rebuild_knowledge_index()
        return f"✅ Index knowledge base dibangun ulang. Dokumen: {docs}, chunks terindex: {chunks}."

    if lower.startswith(("/rag cari ", "/kb cari ")):
        query = raw.split(" ", 2)[2].strip()
        docs = store.search_documents(query, limit=6)
        if not docs:
            return "Belum ada potongan knowledge base yang cocok."
        lines = ["🔎 Hasil Knowledge Base:"]
        for idx, doc in enumerate(docs, start=1):
            snippet = re.sub(r"\s+", " ", doc.get("content", "")).strip()[:420]
            citation = doc.get("citation") or _format_kb_citation(doc)
            lines.append(f"{idx}. {citation} | koleksi {doc.get('collection') or 'Default'} | score {doc.get('score')}\n{snippet}")
        return "\n\n".join(lines)

    if lower.startswith(("/rag tambah", "/kb tambah")):
        body = raw.split(" ", 2)[2].strip() if len(raw.split(" ", 2)) >= 3 else ""
        if "\n" in body:
            title, content = body.split("\n", 1)
        else:
            title, content = "Catatan manual", body
        doc_id, chunks = store.add_document(title=title.strip() or "Catatan manual", text=content, source=f"telegram:{user_id}", collection="Telegram")
        if not chunks:
            return "Gagal menambahkan RAG: isi dokumen kosong."
        return f"✅ Knowledge base ditambahkan. Doc ID: {doc_id}, chunks: {chunks}."


    if lower in {"/gap list", "/gaps", "/knowledge gap"}:
        gaps = store.list_knowledge_gaps(status="open", limit=10)
        if not gaps:
            return "Belum ada knowledge gap terbuka."
        lines = ["🧩 Knowledge gap terbuka:", ""]
        for item in gaps:
            lines.append(f"ID {item.get('id')} | {item.get('intent')} | prioritas {item.get('priority')}\n{item.get('question')}\n")
        return "\n".join(lines).strip()

    if lower.startswith("/gap selesai "):
        gap_id = raw.split(" ", 2)[2].strip()
        ok = store.update_knowledge_gap_status(int(gap_id), status="done") if gap_id.isdigit() else False
        return "✅ Knowledge gap ditandai selesai." if ok else "Gap ID tidak ditemukan."

    if lower in {"/feedback statistik", "/feedback stats"}:
        data = store.feedback_summary(days=30)
        return f"📊 Feedback 30 hari\nTotal: {data.get('total')}\nPositif: {data.get('positive')}\nNegatif: {data.get('negative')}"

    if lower in {"/biaya", "/usage", "/biaya hari ini", "/usage hari ini"}:
        data = store.usage_summary(days=1)
        lines = [
            "📊 Usage 24 jam terakhir",
            "",
            f"Request: {data['requests']}",
            f"Estimasi biaya: Rp{data['cost_idr']:.2f}",
        ]
        rows = data.get("by_model", [])[:10]
        if rows:
            lines.append("\nPer model:")
        for idx, row in enumerate(rows, start=1):
            lines.append(
                f"{idx}. {row.get('model') or '-'}\n"
                f"   Req {row.get('requests')} | In {int(row.get('input_tokens') or 0)} | Out {int(row.get('output_tokens') or 0)} | Rp{float(row.get('cost_idr') or 0):.2f}"
            )
        return "\n".join(lines)

    if lower in {"/model skor", "/model score"}:
        rows = store.model_score_rows(limit=15)
        if not rows:
            return "Belum ada data skor model. Gunakan AI beberapa kali atau jalankan benchmark dulu."
        lines = ["🏆 Skor model adaptif", ""]
        for idx, row in enumerate(rows[:15], start=1):
            lines.append(
                f"{idx}. {row.get('model')}\n"
                f"   Intent: {row.get('intent')} | Score: {row.get('computed_score')}\n"
                f"   Success: {row.get('success_rate')} | Latency: {row.get('avg_latency')}s | Quality: {row.get('avg_quality')}"
            )
        return "\n".join(lines)

    if lower in {"/circuit", "/circuit status"}:
        rows = store.circuit_breaker_status(limit=20)
        if not rows:
            return "Circuit breaker masih kosong."
        lines = ["🧯 Circuit breaker", ""]
        for idx, row in enumerate(rows[:20], start=1):
            status = "BLOCKED" if row.get("blocked") else "OK"
            until = row.get("open_until_wib") or "-"
            lines.append(
                f"{idx}. {row.get('model')}\n"
                f"   Status: {status} | Gagal: {row.get('failure_count')} | Buka sampai: {until}"
            )
        return "\n".join(lines)

    if lower in {"/template", "/template list", "/mode list"}:
        return "Mode tersedia: " + ", ".join(sorted(PROMPT_TEMPLATES.keys()))

    return ""


def run_model_benchmark(
    *,
    store: PowerStore,
    api_url: str,
    api_key: str,
    models: List[str],
    system_prompt: str = "Jawab ringkas dan akurat dalam bahasa Indonesia.",
    timeout: int = 45,
    max_models: int = 8,
) -> List[Dict[str, Any]]:
    tests = [
        ("quick_chat", "Jawab satu kalimat: apa fungsi utama router model AI?"),
        ("deep_reasoning", "Berikan 3 langkah prioritas memperbaiki aplikasi chatbot yang lambat dan sering fallback."),
        ("coding", "Sebutkan bug umum pada Python f-string jika tanda kutip nested salah, lalu berikan contoh perbaikannya."),
        ("academic", "Buat 3 poin latar belakang akademik singkat tentang pentingnya akurasi klaim BPJS."),
    ]
    ranked = store.rank_models_for_intent(list(dict.fromkeys([m for m in models if m])), "general")
    results: List[Dict[str, Any]] = []
    for model in ranked[: max(1, int(max_models or 8))]:
        for task, prompt in tests:
            started = time.time()
            try:
                answer, meta = call_api_once(
                    api_url=api_url,
                    api_key=api_key,
                    model=model,
                    messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": prompt}],
                    temperature=0,
                    max_completion_tokens=900,
                    timeout=timeout,
                )
                latency = round(time.time() - started, 3)
                score = 0.0
                if answer and len(answer.split()) >= 8:
                    score += 0.40
                if any(x in answer.lower() for x in ["router", "model", "fallback", "f-string", "kutip", "langkah", "klaim", "akurasi"]):
                    score += 0.35
                if len(answer) < 1800:
                    score += 0.10
                if latency <= 12:
                    score += 0.15
                score = round(min(score, 1.0), 3)
                store.add_benchmark(model, task, score, latency, True, meta=meta)
                results.append({"model": model, "task": task, "score": score, "latency_seconds": latency, "success": True, "error": ""})
            except Exception as exc:
                latency = round(time.time() - started, 3)
                store.add_benchmark(model, task, 0.0, latency, False, error=str(exc))
                results.append({"model": model, "task": task, "score": 0.0, "latency_seconds": latency, "success": False, "error": str(exc)[:260]})
    return results
